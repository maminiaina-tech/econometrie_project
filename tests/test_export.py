"""Tests for the export module (exporters.py)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

import numpy as np
import openpyxl
import pandas as pd
import pytest
import statsmodels.api as sm
from docx import Document
from statsmodels.regression.linear_model import RegressionResultsWrapper

from econometrie_m1.export.exporters import (
    export_results_to_excel,
    export_results_to_pdf,
    export_results_to_word,
)


@pytest.fixture(scope="module")
def ols_model() -> tuple[RegressionResultsWrapper, pd.DataFrame, pd.DataFrame, pd.Series]:
    """Create a synthetic OLS model for use across all export tests."""
    np.random.seed(42)
    n = 50
    X1 = np.random.randn(n)
    X2 = np.random.randn(n)
    y = 2 + 3 * X1 + 1.5 * X2 + np.random.randn(n) * 0.5
    df = pd.DataFrame({"y": y, "X1": X1, "X2": X2})
    X = sm.add_constant(df[["X1", "X2"]])
    model = sm.OLS(df["y"], X).fit()
    return model, df, X, df["y"]


def _get_model_and_data(
    ols_model: tuple[Any, Any, Any, Any],
) -> tuple[Any, pd.DataFrame, pd.DataFrame, pd.Series]:
    """Unpack the ols_model fixture tuple."""
    model, df, X, y = ols_model
    return model, df, X, y


class TestExportResultsToExcel:
    """Tests for export_results_to_excel."""

    def test_creates_valid_xlsx_file(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Export should produce a readable .xlsx file."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y)

        path = Path(file_path)
        assert path.exists()
        assert path.stat().st_size > 0
        wb = openpyxl.load_workbook(file_path, read_only=True)
        assert wb.sheetnames
        wb.close()

    def test_default_sheets_present(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Default export should include Données, Résumé, Coefficients and Multicolinearite sheets."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y)

        wb = openpyxl.load_workbook(file_path, read_only=True)
        assert "Données" in wb.sheetnames
        assert "Résumé" in wb.sheetnames
        assert "Coefficients" in wb.sheetnames
        assert "Multicolinearite" in wb.sheetnames
        wb.close()

    def test_data_sheet_contains_all_rows(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Data sheet should contain all observations."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y)

        sheet_df = pd.read_excel(file_path, sheet_name="Données")
        assert len(sheet_df) == len(df)
        assert set(df.columns).issubset(set(sheet_df.columns))

    def test_summary_sheet_has_expected_rows(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Summary sheet should contain 9 statistic rows."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y)

        sheet_df = pd.read_excel(file_path, sheet_name="Résumé")
        assert len(sheet_df) == 9
        assert "R²" in sheet_df["Statistique"].values

    def test_coefficients_sheet_has_correct_columns(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Coefficients sheet should have expected columns and one row per variable."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y)

        sheet_df = pd.read_excel(file_path, sheet_name="Coefficients")
        expected_cols = {"Variable", "Coefficient", "Ecart-type", "t", "P>|t|", "Significatif"}
        assert expected_cols.issubset(set(sheet_df.columns))
        assert len(sheet_df) == len(model.params)

    def test_include_matrices_sheets(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """When include_matrices=True, matrix sheets should appear."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y, include_matrices=True)

        wb = openpyxl.load_workbook(file_path, read_only=True)
        assert "Matrice_X" in wb.sheetnames
        assert "Matrice_XtX" in wb.sheetnames
        assert "Matrice_XtY" in wb.sheetnames
        assert "Matrice_XtX_inv" in wb.sheetnames
        assert "Matrice_VarCov" in wb.sheetnames
        wb.close()

    def test_include_diagnostics_sheets(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """When include_diagnostics=True, diagnostic sheets should appear."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y, include_diagnostics=True)

        wb = openpyxl.load_workbook(file_path, read_only=True)
        assert "Tests_Diagnostic" in wb.sheetnames
        assert "Résidus" in wb.sheetnames
        wb.close()

    def test_no_data_sheet_when_excluded(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """When include_data=False, the Données sheet should not be created."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.xlsx")

        export_results_to_excel(file_path, model, df, X, y, include_data=False)

        wb = openpyxl.load_workbook(file_path, read_only=True)
        assert "Données" not in wb.sheetnames
        wb.close()


class TestExportResultsToWord:
    """Tests for export_results_to_word."""

    def test_creates_valid_docx_file(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Export should produce a readable .docx file."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        path = Path(file_path)
        assert path.exists()
        assert path.stat().st_size > 0
        doc = Document(file_path)
        assert len(doc.paragraphs) > 0

    def test_contains_main_heading(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Document should start with the main heading."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        doc = Document(file_path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Résultats Économétriques" in all_text

    def test_contains_data_section(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """With include_data=True, observation count should appear."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        doc = Document(file_path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert f"Nombre d'observations: {len(df)}" in all_text

    def test_contains_summary_table(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """With include_summary=True, a summary table with R² should exist."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        doc = Document(file_path)
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                table_texts.append(row_text)
        combined = "\n".join(table_texts)
        assert "R²" in combined

    def test_contains_coefficients_table(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Coefficient table should list all variable names."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        doc = Document(file_path)
        all_table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_table_text += cell.text + " "
        for var_name in model.params.index:
            assert var_name in all_table_text

    def test_contains_vif_table(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """VIF table should be present with variable names."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y)

        doc = Document(file_path)
        all_table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_table_text += cell.text + " "
        assert "VIF" in all_table_text

    def test_diagnostics_section_when_enabled(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """With include_diagnostics=True, diagnostic test names should appear."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y, include_diagnostics=True)

        doc = Document(file_path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Durbin-Watson" in all_text
        assert "Breusch-Pagan" in all_text

    def test_no_data_section_when_excluded(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """With include_data=False, observation count should not appear."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.docx")

        export_results_to_word(file_path, model, df, X, y, include_data=False)

        doc = Document(file_path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert f"Nombre d'observations: {len(df)}" not in all_text


class TestExportResultsToPdf:
    """Tests for export_results_to_pdf."""

    def test_creates_valid_pdf_file(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Export should produce a non-empty .pdf file."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.pdf")

        export_results_to_pdf(file_path, model, df, X, y)

        path = Path(file_path)
        assert path.exists()
        assert path.stat().st_size > 0

    def test_file_starts_with_pdf_header(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """Output file should start with the %PDF magic bytes."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.pdf")

        export_results_to_pdf(file_path, model, df, X, y)

        with open(file_path, "rb") as f:
            header = f.read(5)
        assert header == b"%PDF-"

    def test_contains_content_when_included(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """PDF with defaults should have substantial content (summary + data + VIF)."""
        model, df, X, y = _get_model_and_data(ols_model)
        file_path = str(tmp_path / "results.pdf")

        export_results_to_pdf(file_path, model, df, X, y)

        path = Path(file_path)
        size_no_diag = path.stat().st_size
        assert size_no_diag > 1000

    def test_diagnostics_increases_file_size(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """PDF with diagnostics enabled should be larger than without."""
        model, df, X, y = _get_model_and_data(ols_model)

        file_base = str(tmp_path / "base.pdf")
        export_results_to_pdf(file_base, model, df, X, y)

        file_diag = str(tmp_path / "diag.pdf")
        export_results_to_pdf(file_diag, model, df, X, y, include_diagnostics=True)

        assert Path(file_diag).stat().st_size > Path(file_base).stat().st_size

    def test_no_data_smaller_file(
        self,
        tmp_path: Path,
        ols_model: tuple[Any, Any, Any, Any],
    ) -> None:
        """PDF without data should be smaller than with data."""
        model, df, X, y = _get_model_and_data(ols_model)

        file_full = str(tmp_path / "full.pdf")
        export_results_to_pdf(file_full, model, df, X, y)

        file_no_data = str(tmp_path / "nodata.pdf")
        export_results_to_pdf(file_no_data, model, df, X, y, include_data=False)

        assert Path(file_no_data).stat().st_size < Path(file_full).stat().st_size
