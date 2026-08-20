"""Application principale d'analyse econometrique avancee."""

from __future__ import annotations

import math
import os
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, messagebox, simpledialog, ttk

import docx
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import statsmodels.api as sm
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from scipy import stats
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from statsmodels.stats.diagnostic import het_breuschpagan, normal_ad
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.stattools import durbin_watson, jarque_bera


class EconometrieAvanceeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Économétrie Avancée - Analyse Complète")
        self.root.geometry("1300x900")

        # Configuration
        self.decimal_places = 4
        self.show_formulas = True
        self.show_interpretations = True
        self.show_calculation_steps = True
        self.show_hypothesis = True
        self.use_alpha = True  # Utiliser a au lieu de beta

        # Variables
        self.data = None
        self.model = None
        self.X = None
        self.y = None
        self.partial_corr_results = None
        self.klein_test_results = None
        self.fg_test_results = None
        self.outliers_indices = []
        self.alpha = 0.05  # Seuil par défaut

        # Interface
        self.create_widgets()
        self.create_menu()

    def set_alpha_level(self):
        try:
            alpha = simpledialog.askfloat(
                "Niveau de significativité",
                "Entrez le seuil alpha (ex: 0.05 pour 5%):",
                parent=self.root,
                minvalue=0.001,
                maxvalue=0.5,
            )
            if alpha is not None:
                self.alpha = alpha
                if self.model:
                    self.show_results()
        except ValueError:
            messagebox.showerror("Erreur", "Veuillez entrer un nombre valide entre 0.001 et 0.5")

    def show_student_table(self):
        """Affiche la table de Student exactement comme dans le PDF fourni"""
        table_window = tk.Toplevel(self.root)
        table_window.title("Table de Student (t) - Bilatéral")
        table_window.geometry("1000x700")

        # Style
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Helvetica", 10, "bold"))
        style.configure("Title.TLabel", font=("Helvetica", 14, "bold"))

        # Titre
        title = ttk.Label(
            table_window,
            text="Table de Student (distribution t)\nAire dans la queue supérieure de la distribution",
            style="Title.TLabel",
        )
        title.pack(pady=10)

        # Cadre pour la table
        frame = ttk.Frame(table_window)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Création de la table
        columns = ("dl", "0.20", "0.10", "0.05", "0.025", "0.01", "0.005")
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=25)

        # Configuration des colonnes
        column_widths = {
            "dl": 100,
            "0.20": 80,
            "0.10": 80,
            "0.05": 80,
            "0.025": 80,
            "0.01": 80,
            "0.005": 80,
        }

        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=column_widths.get(col, 80), anchor=tk.CENTER)

        # Données EXACTES du PDF (corrigées selon vos indications)
        student_data = [
            (1, 1.376, 3.078, 6.314, 12.706, 31.821, 63.656),
            (2, 1.061, 1.886, 2.920, 4.303, 6.965, 9.925),
            (3, 0.978, 1.638, 2.353, 3.182, 4.541, 5.841),
            (4, 0.941, 1.533, 2.132, 2.776, 3.747, 4.604),
            (5, 0.920, 1.476, 2.015, 2.571, 3.365, 4.032),
            (6, 0.906, 1.440, 1.943, 2.447, 3.143, 3.707),
            (7, 0.896, 1.415, 1.895, 2.365, 2.998, 3.499),
            (8, 0.889, 1.397, 1.860, 2.306, 2.896, 3.355),
            (9, 0.883, 1.383, 1.833, 2.262, 2.821, 3.250),
            (10, 0.879, 1.372, 1.812, 2.228, 2.764, 3.169),
            (11, 0.876, 1.363, 1.796, 2.201, 2.718, 3.106),
            (12, 0.873, 1.356, 1.782, 2.179, 2.681, 3.055),
            (13, 0.870, 1.350, 1.771, 2.160, 2.650, 3.012),
            (14, 0.868, 1.345, 1.761, 2.145, 2.624, 2.977),
            (15, 0.866, 1.341, 1.753, 2.131, 2.602, 2.947),
            (16, 0.865, 1.337, 1.746, 2.120, 2.583, 2.921),
            (17, 0.863, 1.333, 1.740, 2.110, 2.567, 2.898),
            (18, 0.862, 1.330, 1.734, 2.101, 2.552, 2.878),
            (19, 0.861, 1.328, 1.729, 2.093, 2.539, 2.861),
            (20, 0.860, 1.325, 1.725, 2.086, 2.528, 2.845),
            (21, 0.859, 1.323, 1.721, 2.080, 2.518, 2.831),
            (22, 0.858, 1.321, 1.717, 2.074, 2.508, 2.819),
            (23, 0.858, 1.319, 1.714, 2.069, 2.500, 2.807),
            (24, 0.857, 1.318, 1.711, 2.064, 2.492, 2.797),
            (25, 0.856, 1.316, 1.708, 2.060, 2.485, 2.787),
            (26, 0.856, 1.315, 1.706, 2.056, 2.479, 2.779),
            (27, 0.855, 1.314, 1.703, 2.052, 2.473, 2.771),
            (28, 0.855, 1.313, 1.701, 2.048, 2.467, 2.763),
            (29, 0.854, 1.311, 1.699, 2.045, 2.462, 2.756),
            (30, 0.854, 1.310, 1.697, 2.042, 2.457, 2.750),
            (31, 0.853, 1.309, 1.696, 2.040, 2.453, 2.744),
            (32, 0.853, 1.309, 1.694, 2.037, 2.449, 2.738),
            (33, 0.853, 1.308, 1.692, 2.035, 2.445, 2.733),
            (34, 0.852, 1.307, 1.691, 2.032, 2.441, 2.728),
            (35, 0.852, 1.306, 1.690, 2.030, 2.438, 2.724),
            (36, 0.852, 1.306, 1.688, 2.028, 2.434, 2.719),
            (37, 0.851, 1.305, 1.687, 2.026, 2.431, 2.715),
            (38, 0.851, 1.304, 1.686, 2.024, 2.429, 2.712),
            (39, 0.851, 1.304, 1.685, 2.023, 2.426, 2.708),
            (40, 0.851, 1.303, 1.684, 2.021, 2.423, 2.704),
            (41, 0.850, 1.302, 1.682, 2.008, 2.421, 2.701),
            (42, 0.850, 1.302, 1.681, 2.017, 2.418, 2.698),
            (43, 0.850, 1.301, 1.680, 2.015, 2.416, 2.695),
            (44, 0.850, 1.301, 1.679, 2.014, 2.414, 2.692),
            (45, 0.850, 1.300, 1.679, 2.013, 2.412, 2.690),
            (46, 0.849, 1.300, 1.678, 2.012, 2.410, 2.687),
            (47, 0.849, 1.299, 1.677, 2.011, 2.408, 2.685),
            (48, 0.849, 1.299, 1.677, 2.010, 2.407, 2.682),
            (49, 0.849, 1.299, 1.676, 2.009, 2.405, 2.680),
            (50, 0.849, 1.298, 1.675, 2.008, 2.403, 2.678),
            (51, 0.848, 1.298, 1.674, 2.007, 2.402, 2.676),
            (52, 0.848, 1.298, 1.674, 2.006, 2.400, 2.674),
            (53, 0.848, 1.297, 1.673, 2.005, 2.399, 2.672),
            (54, 0.848, 1.297, 1.673, 2.004, 2.397, 2.670),
            (55, 0.847, 1.296, 1.672, 2.003, 2.396, 2.668),
            (56, 0.847, 1.296, 1.672, 2.002, 2.395, 2.667),
            (57, 0.847, 1.295, 1.671, 2.002, 2.394, 2.665),
            (58, 0.847, 1.295, 1.671, 2.001, 2.392, 2.663),
            (59, 0.846, 1.295, 1.670, 2.000, 2.391, 2.662),
            (60, 0.846, 1.294, 1.670, 2.000, 2.390, 2.660),
            (61, 0.846, 1.294, 1.669, 1.999, 2.389, 2.658),
            (62, 0.846, 1.293, 1.669, 1.998, 2.388, 2.657),
            (63, 0.846, 1.293, 1.668, 1.998, 2.387, 2.656),
            (64, 0.845, 1.293, 1.668, 1.997, 2.386, 2.655),
            (65, 0.845, 1.292, 1.667, 1.997, 2.385, 2.654),
            (66, 0.845, 1.292, 1.667, 1.996, 2.384, 2.652),
            (67, 0.845, 1.292, 1.666, 1.995, 2.383, 2.651),
            (68, 0.845, 1.291, 1.666, 1.995, 2.382, 2.650),
            (69, 0.844, 1.291, 1.665, 1.994, 2.382, 2.649),
            (70, 0.844, 1.291, 1.665, 1.994, 2.381, 2.648),
            (71, 0.844, 1.290, 1.664, 1.993, 2.380, 2.647),
            (72, 0.844, 1.290, 1.664, 1.993, 2.379, 2.646),
            (73, 0.844, 1.290, 1.663, 1.992, 2.379, 2.645),
            (74, 0.844, 1.290, 1.663, 1.992, 2.378, 2.644),
            (75, 0.843, 1.289, 1.662, 1.991, 2.377, 2.643),
            (76, 0.843, 1.289, 1.662, 1.991, 2.376, 2.642),
            (77, 0.843, 1.289, 1.661, 1.990, 2.376, 2.641),
            (78, 0.843, 1.288, 1.661, 1.990, 2.375, 2.640),
            (79, 0.843, 1.288, 1.660, 1.989, 2.374, 2.640),
            (80, 0.843, 1.288, 1.660, 1.989, 2.374, 2.639),
            (81, 0.843, 1.288, 1.660, 1.988, 2.373, 2.638),
            (82, 0.842, 1.287, 1.659, 1.988, 2.373, 2.637),
            (83, 0.842, 1.287, 1.659, 1.987, 2.372, 2.636),
            (84, 0.842, 1.287, 1.658, 1.987, 2.372, 2.636),
            (85, 0.842, 1.287, 1.658, 1.986, 2.371, 2.635),
            (86, 0.842, 1.286, 1.657, 1.986, 2.370, 2.634),
            (87, 0.842, 1.286, 1.657, 1.986, 2.370, 2.634),
            (88, 0.842, 1.286, 1.656, 1.985, 2.369, 2.633),
            (89, 0.841, 1.286, 1.656, 1.985, 2.369, 2.632),
            (90, 0.841, 1.286, 1.656, 1.984, 2.368, 2.632),
            (91, 0.841, 1.285, 1.655, 1.984, 2.368, 2.631),
            (92, 0.841, 1.285, 1.655, 1.984, 2.367, 2.630),
            (93, 0.841, 1.285, 1.654, 1.983, 2.367, 2.630),
            (94, 0.841, 1.285, 1.654, 1.983, 2.366, 2.629),
            (95, 0.841, 1.285, 1.654, 1.982, 2.366, 2.629),
            (96, 0.841, 1.284, 1.653, 1.982, 2.365, 2.628),
            (97, 0.840, 1.284, 1.653, 1.982, 2.365, 2.627),
            (98, 0.840, 1.284, 1.653, 1.981, 2.364, 2.627),
            (99, 0.840, 1.284, 1.652, 1.981, 2.364, 2.626),
            (100, 0.840, 1.283, 1.652, 1.980, 2.364, 2.626),
            ("∞", 0.842, 1.282, 1.645, 1.960, 2.326, 2.576),
        ]

        for row in student_data:
            tree.insert("", tk.END, values=row)

        # Barre de défilement
        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscroll=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.pack(fill=tk.BOTH, expand=True)

        # Notes
        notes_frame = ttk.Frame(table_window)
        notes_frame.pack(fill=tk.X, padx=10, pady=5)

        note1 = tk.Label(
            notes_frame,
            text="Lecture : Pour dl=30 et α=0.05 (bilatéral), la valeur critique est t=2.042",
            font=("Helvetica", 9),
            anchor="w",
        )
        note1.pack(fill=tk.X)

        note2 = tk.Label(
            notes_frame,
            text="z = valeur limite totale pour α=0.05",
            font=("Helvetica", 9),
            anchor="w",
        )
        note2.pack(fill=tk.X)

    def show_chi2_table(self):
        """Affiche la table complète du Khi-deux de 1 à 30 comme dans le fichier"""
        table_window = tk.Toplevel(self.root)
        table_window.title("Table de Khi-deux (χ²)")
        table_window.geometry("1200x600")

        # Style
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Helvetica", 10, "bold"))

        # Titre
        title = tk.Label(
            table_window,
            text="Table de Khi-deux (distribution χ²)\nP(χ² ≥ χ²v,α) = α",
            font=("Helvetica", 14, "bold"),
        )
        title.pack(pady=10)

        # Cadre pour la table
        frame = ttk.Frame(table_window)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Création de la table
        columns = (
            "v",
            "0.999",
            "0.995",
            "0.99",
            "0.975",
            "0.95",
            "0.90",
            "0.50",
            "0.10",
            "0.05",
            "0.025",
            "0.01",
            "0.005",
            "0.001",
        )
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=15)

        # Configuration des colonnes
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=80, anchor=tk.CENTER)
        tree.column("v", width=50, anchor=tk.CENTER)

        # Données exactes de la table du Khi-deux (1-30)
        chi2_data = [
            (1, 0.00, 0.00, 0.00, 0.00, 0.00, 0.02, 0.45, 2.71, 3.84, 5.02, 6.63, 7.88, 10.83),
            (2, 0.00, 0.01, 0.02, 0.05, 0.10, 0.21, 1.39, 4.61, 5.99, 7.38, 9.21, 10.60, 13.82),
            (3, 0.02, 0.07, 0.11, 0.22, 0.35, 0.58, 2.37, 6.25, 7.81, 9.35, 11.34, 12.84, 16.27),
            (4, 0.09, 0.21, 0.30, 0.48, 0.71, 1.06, 3.36, 7.78, 9.49, 11.14, 13.28, 14.86, 18.47),
            (5, 0.21, 0.41, 0.55, 0.83, 1.15, 1.61, 4.35, 9.24, 11.07, 12.83, 15.09, 16.75, 20.51),
            (6, 0.38, 0.68, 0.87, 1.24, 1.64, 2.20, 5.35, 10.64, 12.59, 14.45, 16.81, 18.55, 22.46),
            (7, 0.60, 0.99, 1.24, 1.69, 2.17, 2.83, 6.35, 12.02, 14.07, 16.01, 18.48, 20.28, 24.32),
            (8, 0.86, 1.34, 1.65, 2.18, 2.73, 3.49, 7.34, 13.36, 15.51, 17.53, 20.09, 21.95, 26.12),
            (9, 1.15, 1.73, 2.09, 2.70, 3.33, 4.17, 8.34, 14.68, 16.92, 19.02, 21.67, 23.59, 27.88),
            (
                10,
                1.48,
                2.16,
                2.56,
                3.25,
                3.94,
                4.87,
                9.34,
                15.99,
                18.31,
                20.48,
                23.21,
                25.19,
                29.59,
            ),
            (
                11,
                1.83,
                2.60,
                3.06,
                3.82,
                4.57,
                5.58,
                10.34,
                17.28,
                19.68,
                21.92,
                24.73,
                26.76,
                31.26,
            ),
            (
                12,
                2.21,
                3.07,
                3.57,
                4.40,
                5.23,
                6.30,
                11.34,
                18.55,
                21.03,
                23.34,
                26.22,
                28.30,
                32.91,
            ),
            (
                13,
                2.62,
                3.57,
                4.11,
                5.01,
                5.89,
                7.04,
                12.34,
                19.81,
                22.36,
                24.74,
                27.69,
                29.82,
                34.53,
            ),
            (
                14,
                3.04,
                4.07,
                4.66,
                5.63,
                6.57,
                7.79,
                13.34,
                21.06,
                23.68,
                26.12,
                29.14,
                31.32,
                36.12,
            ),
            (
                15,
                3.48,
                4.60,
                5.23,
                6.26,
                7.26,
                8.55,
                14.34,
                22.31,
                25.00,
                27.49,
                30.58,
                32.80,
                37.70,
            ),
            (
                16,
                3.94,
                5.14,
                5.81,
                6.91,
                7.96,
                9.31,
                15.34,
                23.54,
                26.30,
                28.85,
                32.00,
                34.27,
                39.25,
            ),
            (
                17,
                4.42,
                5.70,
                6.41,
                7.66,
                8.67,
                10.09,
                16.34,
                24.77,
                27.59,
                30.19,
                33.41,
                35.72,
                40.79,
            ),
            (
                18,
                4.90,
                6.26,
                7.01,
                8.23,
                9.39,
                10.86,
                17.34,
                25.99,
                28.87,
                31.53,
                34.81,
                37.16,
                42.94,
            ),
            (
                19,
                5.41,
                6.84,
                7.63,
                8.91,
                10.12,
                11.65,
                18.34,
                27.20,
                30.14,
                32.85,
                36.19,
                38.58,
                43.82,
            ),
            (
                20,
                5.92,
                7.43,
                8.26,
                9.59,
                10.85,
                12.44,
                19.34,
                28.41,
                31.41,
                34.17,
                37.57,
                40.00,
                45.31,
            ),
            (
                21,
                6.45,
                8.03,
                8.90,
                10.28,
                11.59,
                13.24,
                20.34,
                29.62,
                32.67,
                35.48,
                38.93,
                41.40,
                46.80,
            ),
            (
                22,
                6.98,
                8.64,
                9.54,
                10.98,
                12.34,
                14.04,
                21.34,
                30.81,
                33.92,
                36.78,
                40.29,
                42.80,
                48.27,
            ),
            (
                23,
                7.53,
                9.26,
                10.20,
                11.69,
                13.09,
                14.85,
                22.34,
                32.01,
                35.17,
                38.08,
                41.64,
                44.18,
                49.73,
            ),
            (
                24,
                8.08,
                9.89,
                10.86,
                12.40,
                13.85,
                15.66,
                23.34,
                33.20,
                36.42,
                39.36,
                42.98,
                45.56,
                51.18,
            ),
            (
                25,
                8.65,
                10.52,
                11.52,
                13.12,
                14.61,
                16.47,
                24.34,
                34.38,
                37.65,
                40.65,
                44.31,
                46.93,
                52.62,
            ),
            (
                26,
                9.22,
                11.16,
                12.20,
                13.84,
                15.38,
                17.29,
                25.34,
                35.56,
                38.89,
                41.92,
                45.64,
                48.29,
                54.05,
            ),
            (
                27,
                9.80,
                11.81,
                12.88,
                14.57,
                16.15,
                18.11,
                26.34,
                36.74,
                40.11,
                43.19,
                46.96,
                49.65,
                55.48,
            ),
            (
                28,
                10.39,
                12.46,
                13.56,
                15.31,
                16.93,
                18.94,
                27.34,
                37.92,
                41.34,
                44.46,
                48.28,
                50.99,
                56.89,
            ),
            (
                29,
                10.99,
                13.12,
                14.26,
                16.05,
                17.71,
                19.77,
                28.34,
                39.09,
                42.56,
                45.72,
                49.59,
                52.34,
                58.30,
            ),
            (
                30,
                11.59,
                13.79,
                14.95,
                16.79,
                18.49,
                20.60,
                29.34,
                40.26,
                43.77,
                46.98,
                50.89,
                53.67,
                59.70,
            ),
        ]

        for row in chi2_data:
            tree.insert("", tk.END, values=row)

        # Barre de défilement
        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscroll=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        tree.pack(fill=tk.BOTH, expand=True)

        # Note
        note = tk.Label(
            table_window,
            text="Pour v > 30, la loi du χ² peut être approximée par la loi normale N(v, √v)",
            font=("Helvetica", 9),
        )
        note.pack(pady=5)

    def show_fisher_table(self):
        """Affiche la table complète de Fisher-Snedecor exactement comme dans les fichiers fournis"""
        table_window = tk.Toplevel(self.root)
        table_window.title("Table de Fisher-Snedecor (F)")
        table_window.geometry("1200x800")

        # Style
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Helvetica", 10, "bold"))

        # Titre
        title = tk.Label(
            table_window,
            text="Table de Fisher-Snedecor (distribution F)\nValeurs de F ayant la probabilité P d'être dépassées (F = s₁²/s₂²)",
            font=("Helvetica", 14, "bold"),
        )
        title.pack(pady=10)

        # Notebook pour les différentes tables v1
        notebook = ttk.Notebook(table_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Données complètes extraites des fichiers
        fisher_tables = {
            1: [
                (1, 161.4, 4052.0),
                (2, 18.51, 98.49),
                (3, 10.13, 34.12),
                (4, 7.71, 21.20),
                (5, 6.61, 16.26),
                (6, 5.99, 13.74),
                (7, 5.59, 12.25),
                (8, 5.32, 11.26),
                (9, 5.12, 10.56),
                (10, 4.96, 10.04),
                (11, 4.84, 9.65),
                (12, 4.75, 9.33),
                (13, 4.67, 9.07),
                (14, 4.60, 8.86),
                (15, 4.54, 8.68),
                (16, 4.49, 8.53),
                (17, 4.45, 8.40),
                (18, 4.41, 8.28),
                (19, 4.38, 8.18),
                (20, 4.35, 8.10),
                (21, 4.32, 8.02),
                (22, 4.30, 7.94),
                (23, 4.28, 7.88),
                (24, 4.26, 7.82),
                (25, 4.24, 7.77),
                (26, 4.22, 7.72),
                (27, 4.21, 7.68),
                (28, 4.20, 7.64),
                (29, 4.18, 7.60),
                (30, 4.17, 7.56),
                (40, 4.08, 7.31),
                (60, 4.00, 7.08),
                (120, 3.92, 6.85),
                ("∞", 3.84, 6.64),
            ],
            2: [
                (1, 199.5, 4999.0),
                (2, 19.00, 99.00),
                (3, 9.55, 30.81),
                (4, 6.94, 18.00),
                (5, 5.79, 13.27),
                (6, 5.14, 10.92),
                (7, 4.74, 9.55),
                (8, 4.46, 8.65),
                (9, 4.26, 8.02),
                (10, 4.10, 7.56),
                (11, 3.98, 7.20),
                (12, 3.88, 6.93),
                (13, 3.80, 6.70),
                (14, 3.74, 6.51),
                (15, 3.68, 6.36),
                (16, 3.63, 6.23),
                (17, 3.59, 6.11),
                (18, 3.55, 6.01),
                (19, 3.52, 5.93),
                (20, 3.49, 5.85),
                (21, 3.47, 5.78),
                (22, 3.44, 5.72),
                (23, 3.42, 5.66),
                (24, 3.40, 5.61),
                (25, 3.38, 5.57),
                (26, 3.37, 5.53),
                (27, 3.35, 5.49),
                (28, 3.34, 5.45),
                (29, 3.33, 5.42),
                (30, 3.32, 5.39),
                (40, 3.23, 5.18),
                (60, 3.15, 4.98),
                (120, 3.07, 4.79),
                ("∞", 2.99, 4.61),
            ],
            3: [
                (1, 215.7, 5403.0),
                (2, 19.16, 99.17),
                (3, 9.28, 29.46),
                (4, 6.59, 16.69),
                (5, 5.41, 12.06),
                (6, 4.76, 9.78),
                (7, 4.35, 8.45),
                (8, 4.07, 7.59),
                (9, 3.86, 6.99),
                (10, 3.71, 6.55),
                (11, 3.59, 6.22),
                (12, 3.49, 5.95),
                (13, 3.41, 5.74),
                (14, 3.34, 5.56),
                (15, 3.29, 5.42),
                (16, 3.24, 5.29),
                (17, 3.20, 5.18),
                (18, 3.16, 5.09),
                (19, 3.13, 5.01),
                (20, 3.10, 4.94),
                (21, 3.07, 4.87),
                (22, 3.05, 4.82),
                (23, 3.03, 4.76),
                (24, 3.01, 4.72),
                (25, 2.99, 4.68),
                (26, 2.98, 4.64),
                (27, 2.96, 4.60),
                (28, 2.95, 4.57),
                (29, 2.93, 4.54),
                (30, 2.92, 4.51),
                (40, 2.84, 4.31),
                (60, 2.76, 4.13),
                (120, 2.68, 3.95),
                ("∞", 2.60, 3.78),
            ],
            4: [
                (1, 224.6, 5625.0),
                (2, 19.25, 99.25),
                (3, 9.12, 28.71),
                (4, 6.39, 15.98),
                (5, 5.19, 11.39),
                (6, 4.53, 9.15),
                (7, 4.12, 7.85),
                (8, 3.84, 7.01),
                (9, 3.63, 6.42),
                (10, 3.48, 5.99),
                (11, 3.36, 5.67),
                (12, 3.26, 5.41),
                (13, 3.18, 5.20),
                (14, 3.11, 5.03),
                (15, 3.06, 4.89),
                (16, 3.01, 4.77),
                (17, 2.96, 4.67),
                (18, 2.93, 4.58),
                (19, 2.90, 4.50),
                (20, 2.87, 4.43),
                (21, 2.84, 4.37),
                (22, 2.82, 4.31),
                (23, 2.80, 4.26),
                (24, 2.78, 4.22),
                (25, 2.76, 4.18),
                (26, 2.74, 4.14),
                (27, 2.73, 4.11),
                (28, 2.71, 4.07),
                (29, 2.70, 4.04),
                (30, 2.69, 4.02),
                (40, 2.61, 3.83),
                (60, 2.52, 3.65),
                (120, 2.45, 3.48),
                ("∞", 2.37, 3.32),
            ],
            5: [
                (1, 230.2, 5764.0),
                (2, 19.30, 99.30),
                (3, 9.01, 28.24),
                (4, 6.26, 15.52),
                (5, 5.05, 10.97),
                (6, 4.39, 8.75),
                (7, 3.97, 7.46),
                (8, 3.69, 6.63),
                (9, 3.48, 6.06),
                (10, 3.33, 5.64),
                (11, 3.20, 5.32),
                (12, 3.11, 5.06),
                (13, 3.02, 4.86),
                (14, 2.96, 4.69),
                (15, 2.90, 4.56),
                (16, 2.85, 4.44),
                (17, 2.81, 4.34),
                (18, 2.77, 4.25),
                (19, 2.74, 4.17),
                (20, 2.71, 4.10),
                (21, 2.68, 4.04),
                (22, 2.66, 3.99),
                (23, 2.64, 3.94),
                (24, 2.62, 3.90),
                (25, 2.60, 3.86),
                (26, 2.59, 3.82),
                (27, 2.57, 3.78),
                (28, 2.56, 3.75),
                (29, 2.55, 3.73),
                (30, 2.53, 3.70),
                (40, 2.45, 3.51),
                (60, 2.37, 3.34),
                (120, 2.29, 3.17),
                ("∞", 2.21, 3.02),
            ],
            6: [
                (1, 234.0, 5859.0),
                (2, 19.33, 99.33),
                (3, 8.94, 27.91),
                (4, 6.16, 15.21),
                (5, 4.95, 10.67),
                (6, 4.28, 8.47),
                (7, 3.87, 7.19),
                (8, 3.58, 6.37),
                (9, 3.37, 5.80),
                (10, 3.22, 5.39),
                (11, 3.09, 5.07),
                (12, 3.00, 4.82),
                (13, 2.92, 4.62),
                (14, 2.85, 4.46),
                (15, 2.79, 4.32),
                (16, 2.74, 4.20),
                (17, 2.70, 4.10),
                (18, 2.66, 4.01),
                (19, 2.63, 3.94),
                (20, 2.60, 3.87),
                (21, 2.57, 3.81),
                (22, 2.55, 3.76),
                (23, 2.53, 3.71),
                (24, 2.51, 3.67),
                (25, 2.49, 3.63),
                (26, 2.47, 3.59),
                (27, 2.46, 3.56),
                (28, 2.44, 3.53),
                (29, 2.43, 3.50),
                (30, 2.42, 3.47),
                (40, 2.34, 3.29),
                (60, 2.25, 3.12),
                (120, 2.17, 2.96),
                ("∞", 2.10, 2.80),
            ],
            8: [
                (1, 238.9, 5981.0),
                (2, 19.37, 99.37),
                (3, 8.84, 27.49),
                (4, 6.04, 14.80),
                (5, 4.82, 10.27),
                (6, 4.15, 8.10),
                (7, 3.73, 6.84),
                (8, 3.44, 6.03),
                (9, 3.23, 5.47),
                (10, 3.07, 5.06),
                (11, 2.95, 4.74),
                (12, 2.85, 4.50),
                (13, 2.77, 4.30),
                (14, 2.70, 4.14),
                (15, 2.64, 4.00),
                (16, 2.59, 3.89),
                (17, 2.55, 3.79),
                (18, 2.51, 3.71),
                (19, 2.48, 3.63),
                (20, 2.45, 3.56),
                (21, 2.42, 3.51),
                (22, 2.40, 3.45),
                (23, 2.38, 3.41),
                (24, 2.36, 3.36),
                (25, 2.34, 3.32),
                (26, 2.32, 3.29),
                (27, 2.30, 3.26),
                (28, 2.29, 3.23),
                (29, 2.28, 3.20),
                (30, 2.27, 3.17),
                (40, 2.18, 2.99),
                (60, 2.10, 2.82),
                (120, 2.01, 2.66),
                ("∞", 1.94, 2.51),
            ],
            12: [
                (1, 243.9, 6106.0),
                (2, 19.41, 99.42),
                (3, 8.74, 27.05),
                (4, 5.91, 14.37),
                (5, 4.68, 9.89),
                (6, 4.00, 7.72),
                (7, 3.57, 6.47),
                (8, 3.28, 5.67),
                (9, 3.07, 5.11),
                (10, 2.91, 4.71),
                (11, 2.79, 4.40),
                (12, 2.69, 4.16),
                (13, 2.60, 3.96),
                (14, 2.53, 3.80),
                (15, 2.48, 3.67),
                (16, 2.42, 3.55),
                (17, 2.38, 3.45),
                (18, 2.34, 3.37),
                (19, 2.31, 3.30),
                (20, 2.28, 3.23),
                (21, 2.25, 3.17),
                (22, 2.23, 3.12),
                (23, 2.20, 3.07),
                (24, 2.18, 3.03),
                (25, 2.16, 2.99),
                (26, 2.15, 2.96),
                (27, 2.13, 2.93),
                (28, 2.12, 2.90),
                (29, 2.10, 2.87),
                (30, 2.09, 2.84),
                (40, 2.00, 2.66),
                (60, 1.92, 2.50),
                (120, 1.83, 2.34),
                ("∞", 1.75, 2.18),
            ],
            24: [
                (1, 249.0, 6234.0),
                (2, 19.45, 99.46),
                (3, 8.64, 26.60),
                (4, 5.77, 13.93),
                (5, 4.53, 9.47),
                (6, 3.84, 7.31),
                (7, 3.41, 6.07),
                (8, 3.12, 5.28),
                (9, 2.90, 4.73),
                (10, 2.74, 4.33),
                (11, 2.61, 4.02),
                (12, 2.50, 3.78),
                (13, 2.42, 3.59),
                (14, 2.35, 3.43),
                (15, 2.29, 3.29),
                (16, 2.24, 3.18),
                (17, 2.19, 3.08),
                (18, 2.15, 3.00),
                (19, 2.11, 2.92),
                (20, 2.08, 2.86),
                (21, 2.05, 2.80),
                (22, 2.03, 2.75),
                (23, 2.01, 2.70),
                (24, 1.98, 2.66),
                (25, 1.96, 2.62),
                (26, 1.95, 2.58),
                (27, 1.93, 2.55),
                (28, 1.91, 2.52),
                (29, 1.90, 2.49),
                (30, 1.89, 2.47),
                (40, 1.79, 2.29),
                (60, 1.70, 2.12),
                (120, 1.61, 1.95),
                ("∞", 1.52, 1.79),
            ],
            "∞": [
                (1, 254.3, 6366.0),
                (2, 19.50, 99.50),
                (3, 8.53, 26.12),
                (4, 5.63, 13.46),
                (5, 4.36, 9.02),
                (6, 3.67, 6.88),
                (7, 3.23, 5.65),
                (8, 2.93, 4.86),
                (9, 2.71, 4.31),
                (10, 2.54, 3.91),
                (11, 2.40, 3.60),
                (12, 2.30, 3.36),
                (13, 2.21, 3.16),
                (14, 2.13, 3.00),
                (15, 2.07, 2.87),
                (16, 2.01, 2.75),
                (17, 1.96, 2.65),
                (18, 1.92, 2.57),
                (19, 1.88, 2.49),
                (20, 1.84, 2.42),
                (21, 1.81, 2.36),
                (22, 1.78, 2.31),
                (23, 1.76, 2.26),
                (24, 1.73, 2.21),
                (25, 1.71, 2.17),
                (26, 1.69, 2.13),
                (27, 1.67, 2.10),
                (28, 1.65, 2.06),
                (29, 1.64, 2.03),
                (30, 1.62, 2.01),
                (40, 1.51, 1.80),
                (60, 1.39, 1.60),
                (120, 1.25, 1.38),
                ("∞", 1.00, 1.00),
            ],
        }

        # Création des onglets avec les données complètes
        for v1 in [1, 2, 3, 4, 5, 6, 8, 12, 24, "∞"]:
            tab = ttk.Frame(notebook)
            notebook.add(tab, text=f"v1 = {v1}")

            frame = ttk.Frame(tab)
            frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

            columns = ("v2", "P=0.05", "P=0.01")
            tree = ttk.Treeview(frame, columns=columns, show="headings", height=20)

            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=100, anchor=tk.CENTER)
            tree.column("v2", width=80, anchor=tk.CENTER)

            for row in fisher_tables[v1]:
                tree.insert("", tk.END, values=row)

            scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tree.yview)
            tree.configure(yscroll=scrollbar.set)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            tree.pack(fill=tk.BOTH, expand=True)

        # Note
        note = tk.Label(
            table_window,
            text="Note : s₁² est la plus grande des deux variances estimées, avec v1 degrés de liberté",
            font=("Helvetica", 9),
        )
        note.pack(pady=5)

    def create_menu(self):
        menubar = tk.Menu(self.root)

        # Menu Configuration
        config_menu = tk.Menu(menubar, tearoff=0)
        config_menu.add_command(label="Décimales...", command=self.set_decimal_places)
        config_menu.add_command(label="Seuil alpha...", command=self.set_alpha_level)
        config_menu.add_checkbutton(
            label="Afficher les formules",
            variable=tk.BooleanVar(value=self.show_formulas),
            command=self.toggle_formulas,
        )
        config_menu.add_checkbutton(
            label="Afficher les interprétations",
            variable=tk.BooleanVar(value=self.show_interpretations),
            command=self.toggle_interpretations,
        )
        config_menu.add_checkbutton(
            label="Afficher étapes calcul",
            variable=tk.BooleanVar(value=self.show_calculation_steps),
            command=self.toggle_calculation_steps,
        )
        config_menu.add_checkbutton(
            label="Afficher hypothèses tests",
            variable=tk.BooleanVar(value=self.show_hypothesis),
            command=self.toggle_hypothesis,
        )
        config_menu.add_checkbutton(
            label="Utiliser a au lieu de beta",
            variable=tk.BooleanVar(value=self.use_alpha),
            command=self.toggle_alpha_beta,
        )
        menubar.add_cascade(label="Configuration", menu=config_menu)

        # Menu Tables
        tables_menu = tk.Menu(menubar, tearoff=0)
        tables_menu.add_command(label="Table de Student (t)", command=self.show_student_table)
        tables_menu.add_command(label="Table de Fisher (F)", command=self.show_fisher_table)
        tables_menu.add_command(label="Table de Khi-2 (χ²)", command=self.show_chi2_table)
        menubar.add_cascade(label="Tables Statistiques", menu=tables_menu)

        self.root.config(menu=menubar)

    def toggle_alpha_beta(self):
        self.use_alpha = not self.use_alpha
        if self.model:
            self.show_results()

    def set_decimal_places(self):
        try:
            places = simpledialog.askinteger(
                "Décimales",
                "Nombre de décimales à afficher:",
                parent=self.root,
                minvalue=0,
                maxvalue=10,
            )
            if places is not None:
                self.decimal_places = places
                if self.model:
                    self.show_results()
        except ValueError:
            messagebox.showerror("Erreur", "Veuillez entrer un nombre valide (0-10)")

    def toggle_formulas(self):
        self.show_formulas = not self.show_formulas
        if self.model:
            self.show_results()

    def toggle_interpretations(self):
        self.show_interpretations = not self.show_interpretations
        if self.model:
            self.show_results()

    def toggle_calculation_steps(self):
        self.show_calculation_steps = not self.show_calculation_steps
        if self.model:
            self.show_results()

    def toggle_hypothesis(self):
        self.show_hypothesis = not self.show_hypothesis
        if self.model:
            self.show_results()

    def create_widgets(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(expand=True, fill="both")

        # Onglet Données
        self.data_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.data_tab, text="Données")
        self.setup_data_tab()

        # Onglet Modélisation
        self.model_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.model_tab, text="Modélisation")
        self.setup_model_tab()

        # Onglet Résultats
        self.results_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.results_tab, text="Résultats Complets")
        self.setup_results_tab()

        # Onglet Diagnostic
        self.diag_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.diag_tab, text="Diagnostics")
        self.setup_diag_tab()

        # Onglet Multicolinéarité
        self.multicollinearity_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.multicollinearity_tab, text="Multicolinéarité")
        self.setup_multicollinearity_tab()

        # Onglet Corrélation Partielle
        self.partial_corr_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.partial_corr_tab, text="Corrélation Partielle")
        self.setup_partial_corr_tab()

        # Onglet Tests Statistiques
        self.stat_tests_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.stat_tests_tab, text="Tests Statistiques")
        self.setup_stat_tests_tab()

        # Onglet Export
        self.export_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.export_tab, text="Export")
        self.setup_export_tab()

    def setup_data_tab(self):
        frame = ttk.Frame(self.data_tab)
        frame.pack(pady=10)

        ttk.Button(frame, text="Importer Excel", command=self.import_data).pack(side="left", padx=5)
        ttk.Button(frame, text="Afficher Stats", command=self.show_stats).pack(side="left", padx=5)
        ttk.Button(frame, text="Nettoyer Données", command=self.clean_data).pack(
            side="left", padx=5
        )
        ttk.Button(frame, text="Détecter Valeurs Aberrantes", command=self.detect_outliers).pack(
            side="left", padx=5
        )
        ttk.Button(frame, text="Corriger Valeurs Aberrantes", command=self.correct_outliers).pack(
            side="left", padx=5
        )

        self.data_tree = ttk.Treeview(self.data_tab)
        self.data_tree.pack(expand=True, fill="both", padx=10, pady=10)

        scrollbar = ttk.Scrollbar(self.data_tree, orient="vertical", command=self.data_tree.yview)
        scrollbar.pack(side="right", fill="y")
        self.data_tree.configure(yscrollcommand=scrollbar.set)

    def setup_model_tab(self):
        frame = ttk.LabelFrame(self.model_tab, text="Spécification du Modèle")
        frame.pack(pady=10, padx=10, fill="x")

        ttk.Label(frame, text="Variable dépendante (Y):").grid(row=0, column=0, padx=5, pady=5)
        self.y_var = tk.StringVar()
        self.y_menu = ttk.Combobox(frame, textvariable=self.y_var)
        self.y_menu.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        ttk.Label(frame, text="Variables indépendantes (X):").grid(row=1, column=0, padx=5, pady=5)
        self.x_vars = tk.StringVar()
        self.x_entry = ttk.Entry(frame, textvariable=self.x_vars)
        self.x_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        ttk.Button(frame, text="Estimer le Modèle", command=self.estimate_model).grid(
            row=2, column=0, columnspan=2, pady=10
        )

        pred_frame = ttk.LabelFrame(self.model_tab, text="Prédiction")
        pred_frame.pack(pady=10, padx=10, fill="x")

        self.pred_entries = {}
        ttk.Button(pred_frame, text="Calculer Prédiction", command=self.calculate_prediction).pack(
            pady=5
        )
        self.pred_result = ttk.Label(pred_frame, text="")
        self.pred_result.pack()

        # Ajout des champs pour la prévision avec nouvelles valeurs
        self.pred_new_frame = ttk.LabelFrame(self.model_tab, text="Nouvelle Prévision")
        self.pred_new_frame.pack(pady=10, padx=10, fill="x")

        # Ces champs seront remplis dynamiquement dans update_prediction_entries()

    def setup_results_tab(self):
        self.results_text = tk.Text(self.results_tab, wrap="word", font=("Courier", 10))
        self.results_text.pack(expand=True, fill="both", padx=10, pady=10)

        scrollbar = ttk.Scrollbar(
            self.results_text, orient="vertical", command=self.results_text.yview
        )
        scrollbar.pack(side="right", fill="y")
        self.results_text.configure(yscrollcommand=scrollbar.set)

    def setup_diag_tab(self):
        self.figure = plt.figure(figsize=(10, 6), dpi=100)
        self.canvas = FigureCanvasTkAgg(self.figure, master=self.diag_tab)
        self.canvas.get_tk_widget().pack(expand=True, fill="both", padx=10, pady=10)

        frame = ttk.Frame(self.diag_tab)
        frame.pack(fill="x")

        self.diag_type = tk.StringVar()
        diag_types = [
            "Résidus vs Ajustés",
            "QQ Plot",
            "Leverage",
            "Histogramme Résidus",
            "Corrélation",
            "Autocorrélation",
            "Autocorrélation Partielle",
        ]
        self.diag_menu = ttk.Combobox(frame, textvariable=self.diag_type, values=diag_types)
        self.diag_menu.pack(side="left", padx=5)
        self.diag_menu.current(0)

        ttk.Button(frame, text="Afficher Diagnostic", command=self.show_diagnostic).pack(
            side="left", padx=5
        )

    def setup_multicollinearity_tab(self):
        frame = ttk.Frame(self.multicollinearity_tab)
        frame.pack(pady=10, fill="x")

        ttk.Button(
            frame, text="Analyser Multicolinéarité", command=self.analyze_multicollinearity
        ).pack()

        self.multicoll_text = tk.Text(self.multicollinearity_tab, wrap="word", font=("Courier", 10))
        self.multicoll_text.pack(expand=True, fill="both", padx=10, pady=10)

        scrollbar = ttk.Scrollbar(
            self.multicoll_text, orient="vertical", command=self.multicoll_text.yview
        )
        scrollbar.pack(side="right", fill="y")
        self.multicoll_text.configure(yscrollcommand=scrollbar.set)

    def setup_partial_corr_tab(self):
        frame = ttk.Frame(self.partial_corr_tab)
        frame.pack(pady=10, fill="x")

        ttk.Button(
            frame,
            text="Calculer Corrélations Partielles",
            command=self.calculate_partial_correlations,
        ).pack()

        self.partial_corr_text = tk.Text(self.partial_corr_tab, wrap="word", font=("Courier", 10))
        self.partial_corr_text.pack(expand=True, fill="both", padx=10, pady=10)

        scrollbar = ttk.Scrollbar(
            self.partial_corr_text, orient="vertical", command=self.partial_corr_text.yview
        )
        scrollbar.pack(side="right", fill="y")
        self.partial_corr_text.configure(yscrollcommand=scrollbar.set)

    def setup_stat_tests_tab(self):
        frame = ttk.Frame(self.stat_tests_tab)
        frame.pack(pady=10, fill="x")

        ttk.Button(
            frame, text="Tester Significativité Variable", command=self.test_variable_significance
        ).pack(side="left", padx=5)
        ttk.Button(
            frame, text="Tester Significativité Globale", command=self.test_global_significance
        ).pack(side="left", padx=5)
        ttk.Button(frame, text="Test de Klein", command=self.klein_test).pack(side="left", padx=5)
        ttk.Button(frame, text="Test de Farrar-Glauber", command=self.farrar_glauber_test).pack(
            side="left", padx=5
        )
        ttk.Button(frame, text="Déterminer MIX optimal", command=self.determine_optimal_mix).pack(
            side="left", padx=5
        )

        self.stat_tests_text = tk.Text(self.stat_tests_tab, wrap="word", font=("Courier", 10))
        self.stat_tests_text.pack(expand=True, fill="both", padx=10, pady=10)

        scrollbar = ttk.Scrollbar(
            self.stat_tests_text, orient="vertical", command=self.stat_tests_text.yview
        )
        scrollbar.pack(side="right", fill="y")
        self.stat_tests_text.configure(yscrollcommand=scrollbar.set)

    def setup_export_tab(self):
        frame = ttk.Frame(self.export_tab)
        frame.pack(pady=20)

        ttk.Button(frame, text="Exporter vers Excel", command=self.export_to_excel).pack(
            pady=5, fill="x"
        )
        # ttk.Button(frame, text="Exporter vers Word", command=self.export_to_word).pack(pady=5, fill='x')
        ttk.Button(frame, text="Exporter vers PDF", command=self.export_to_pdf).pack(
            pady=5, fill="x"
        )
        ttk.Button(frame, text="Exporter Graphiques", command=self.export_plots).pack(
            pady=5, fill="x"
        )

        self.include_data = tk.BooleanVar(value=True)
        self.include_summary = tk.BooleanVar(value=True)
        self.include_matrices = tk.BooleanVar(value=True)
        self.include_diagnostics = tk.BooleanVar(value=True)
        self.include_plots = tk.BooleanVar(value=True)
        self.include_formulas = tk.BooleanVar(value=True)
        self.include_interpretations = tk.BooleanVar(value=True)
        self.include_partial_corr = tk.BooleanVar(value=True)
        self.include_calculation_steps = tk.BooleanVar(value=True)
        self.include_hypothesis = tk.BooleanVar(value=True)
        self.include_stat_tests = tk.BooleanVar(value=True)

        opt_frame = ttk.LabelFrame(self.export_tab, text="Options d'Export")
        opt_frame.pack(pady=10, padx=10, fill="x")

        ttk.Checkbutton(opt_frame, text="Inclure données", variable=self.include_data).pack(
            anchor="w"
        )
        ttk.Checkbutton(
            opt_frame, text="Inclure résumé modèle", variable=self.include_summary
        ).pack(anchor="w")
        ttk.Checkbutton(opt_frame, text="Inclure matrices", variable=self.include_matrices).pack(
            anchor="w"
        )
        ttk.Checkbutton(
            opt_frame, text="Inclure diagnostics", variable=self.include_diagnostics
        ).pack(anchor="w")
        ttk.Checkbutton(opt_frame, text="Inclure graphiques", variable=self.include_plots).pack(
            anchor="w"
        )
        ttk.Checkbutton(opt_frame, text="Inclure formules", variable=self.include_formulas).pack(
            anchor="w"
        )
        ttk.Checkbutton(
            opt_frame, text="Inclure interprétations", variable=self.include_interpretations
        ).pack(anchor="w")
        ttk.Checkbutton(
            opt_frame, text="Inclure corrélations partielles", variable=self.include_partial_corr
        ).pack(anchor="w")
        ttk.Checkbutton(
            opt_frame, text="Inclure étapes calcul", variable=self.include_calculation_steps
        ).pack(anchor="w")
        ttk.Checkbutton(
            opt_frame, text="Inclure hypothèses tests", variable=self.include_hypothesis
        ).pack(anchor="w")
        ttk.Checkbutton(
            opt_frame, text="Inclure tests statistiques", variable=self.include_stat_tests
        ).pack(anchor="w")

    def detect_outliers(self):
        if self.data is None:
            messagebox.showwarning("Avertissement", "Importez d'abord des données!")
            return

        try:
            self.results_text.delete(1.0, tk.END)
            self.results_text.insert(tk.END, "=== DÉTECTION DE VALEURS ABERRANTES ===\n\n")

            self.outliers_indices = []

            for col in self.data.select_dtypes(include=[np.number]).columns:
                # Méthode IQR
                Q1 = self.data[col].quantile(0.25)
                Q3 = self.data[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR

                outliers = self.data[
                    (self.data[col] < lower_bound) | (self.data[col] > upper_bound)
                ]
                self.outliers_indices.extend(outliers.index.tolist())

                self.results_text.insert(tk.END, f"Variable {col}:\n")
                self.results_text.insert(tk.END, f"- Limite inférieure: {lower_bound:.4f}\n")
                self.results_text.insert(tk.END, f"- Limite supérieure: {upper_bound:.4f}\n")
                self.results_text.insert(
                    tk.END, f"- Nombre de valeurs aberrantes: {len(outliers)}\n"
                )

                if len(outliers) > 0:
                    self.results_text.insert(
                        tk.END,
                        "- Indices des valeurs aberrantes: "
                        + ", ".join(map(str, outliers.index.tolist()))
                        + "\n",
                    )

                self.results_text.insert(tk.END, "\n")

            # Enlever les doublons
            self.outliers_indices = list(set(self.outliers_indices))

            if self.show_formulas:
                self.results_text.insert(tk.END, "\n=== MÉTHODE UTILISÉE ===\n")
                self.results_text.insert(
                    tk.END,
                    "Méthode IQR (Interquartile Range):\n"
                    "1. Calculer Q1 (25ème percentile) et Q3 (75ème percentile)\n"
                    "2. Calculer IQR = Q3 - Q1\n"
                    "3. Déterminer les limites:\n"
                    "   - Limite inférieure = Q1 - 1.5 * IQR\n"
                    "   - Limite supérieure = Q3 + 1.5 * IQR\n"
                    "4. Les valeurs en dehors de ces limites sont considérées comme aberrantes\n",
                )

            if self.show_interpretations:
                self.results_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                if len(self.outliers_indices) > 0:
                    self.results_text.insert(
                        tk.END,
                        f"Des valeurs aberrantes ont été détectées aux indices: {', '.join(map(str, self.outliers_indices))}\n"
                        "Ces valeurs peuvent fausser les résultats de la régression.\n"
                        "Vous pouvez les corriger en utilisant le bouton 'Corriger Valeurs Aberrantes'.\n"
                        "Une méthode courante est de remplacer ces valeurs par la médiane ou la moyenne.\n",
                    )
                else:
                    self.results_text.insert(
                        tk.END,
                        "Aucune valeur aberrante significative n'a été détectée.\n"
                        "Les données semblent propres pour l'analyse.\n",
                    )

            messagebox.showinfo("Succès", "Détection des valeurs aberrantes terminée!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur de détection: {str(e)}")

    def correct_outliers(self):
        if len(self.outliers_indices) == 0:
            messagebox.showwarning(
                "Avertissement", "Aucune valeur aberrante détectée ou données non chargées!"
            )
            return

        try:
            # Créer une copie des données originales pour la correction
            corrected_data = self.data.copy()

            for col in corrected_data.select_dtypes(include=[np.number]).columns:
                # Calcul des bornes
                Q1 = corrected_data[col].quantile(0.25)
                Q3 = corrected_data[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR

                # Remplacer les valeurs aberrantes par la médiane
                median_val = corrected_data[col].median()
                corrected_data.loc[
                    (corrected_data[col] < lower_bound) | (corrected_data[col] > upper_bound), col
                ] = median_val

            self.data = corrected_data
            self.display_data()
            self.outliers_indices = []

            self.results_text.insert(tk.END, "\n=== CORRECTION DES VALEURS ABERRANTES ===\n")
            self.results_text.insert(
                tk.END,
                "Les valeurs aberrantes ont été remplacées par la médiane de chaque variable.\n"
                "Les données ont été mises à jour.\n",
            )

            messagebox.showinfo("Succès", "Valeurs aberrantes corrigées avec succès!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la correction: {str(e)}")

    def calculate_partial_correlations(self):
        if self.data is None:
            messagebox.showwarning("Avertissement", "Importez d'abord des données!")
            return

        try:
            y = self.y_var.get()
            x_vars = [x.strip() for x in self.x_vars.get().split("+")]

            if len(x_vars) < 2:
                messagebox.showwarning(
                    "Avertissement",
                    "Il faut au moins 2 variables pour calculer les corrélations partielles!",
                )
                return

            self.partial_corr_results = {}
            df = self.data.copy()

            self.partial_corr_text.delete(1.0, tk.END)
            self.partial_corr_text.insert(tk.END, "=== CORRÉLATIONS PARTIELLES ===\n\n")

            # Fonction pour calculer les corrélations partielles d'ordre n
            def calculate_partial_corr(var1, var2, control_vars):
                if len(control_vars) > 0:
                    res1 = sm.OLS(df[var1], sm.add_constant(df[control_vars])).fit()
                    res2 = sm.OLS(df[var2], sm.add_constant(df[control_vars])).fit()
                    residuals1 = res1.resid
                    residuals2 = res2.resid
                    return np.corrcoef(residuals1, residuals2)[0, 1]
                else:
                    return df[[var1, var2]].corr().iloc[0, 1]

            # Calcul systématique pour tous les ordres possibles
            for order in range(1, len(x_vars)):
                self.partial_corr_text.insert(
                    tk.END, f"\n=== CORRÉLATIONS PARTIELLES D'ORDRE {order} ===\n"
                )

                # Calcul pour toutes les paires de variables
                for i, var1 in enumerate(x_vars):
                    for var2 in x_vars[i + 1 :]:
                        # Variables de contrôle (toutes sauf var1 et var2)
                        all_controls = [v for v in x_vars if v not in [var1, var2]]

                        # Pour l'ordre n, on prend toutes les combinaisons possibles de n variables de contrôle
                        if len(all_controls) >= order:
                            from itertools import combinations

                            for control_vars in combinations(all_controls, order):
                                partial_corr = calculate_partial_corr(
                                    var1, var2, list(control_vars)
                                )

                                # Stockage des résultats
                                key = (var1, var2, *control_vars)
                                self.partial_corr_results[key] = partial_corr

                                # Affichage
                                self.partial_corr_text.insert(
                                    tk.END,
                                    f"r({var1},{var2} | {','.join(control_vars)}) = {partial_corr:.4f}\n",
                                )

            # Calcul des R² partiels (variation expliquée supplémentaire)
            self.partial_corr_text.insert(tk.END, "\n=== R² PARTIELS ===\n")
            for var in x_vars:
                # Modèle complet
                full_model = sm.OLS(df[y], sm.add_constant(df[x_vars])).fit()
                # Modèle sans la variable d'intérêt
                reduced_vars = [v for v in x_vars if v != var]
                reduced_model = sm.OLS(df[y], sm.add_constant(df[reduced_vars])).fit()

                # Calcul du R² partiel
                ss_full = full_model.ess
                ss_reduced = reduced_model.ess
                r2_partial = (ss_full - ss_reduced) / (1 - reduced_model.rsquared)

                self.partial_corr_text.insert(
                    tk.END, f"R² partiel pour {var} (apport marginal) = {r2_partial:.4f}\n"
                )

            # Formules détaillées
            if self.show_formulas:
                self.partial_corr_text.insert(tk.END, "\n=== FORMULES DÉTAILLÉES ===\n")
                self.partial_corr_text.insert(
                    tk.END,
                    "1. Corrélation partielle d'ordre n:\n"
                    "   - Régression de X1 sur Z1..Zn → Résidus e1\n"
                    "   - Régression de X2 sur Z1..Zn → Résidus e2\n"
                    "   - Corrélation(e1, e2)\n\n"
                    "2. R² partiel (coefficient de détermination partiel):\n"
                    "   R²_partiel = (SSE_reduced - SSE_full) / (1 - R²_reduced)\n"
                    "   où SSE = Somme des carrés des erreurs\n\n"
                    "3. Formule matricielle alternative:\n"
                    "   ρ(X,Y|Z) = -pXY / sqrt(pXX * pYY)\n"
                    "   où pij sont les éléments de la matrice de précision (inverse de la matrice de corrélation)\n",
                )

            # Interprétations avancées
            if self.show_interpretations:
                self.partial_corr_text.insert(tk.END, "\n=== INTERPRÉTATIONS AVANCÉES ===\n")
                for key, corr in self.partial_corr_results.items():
                    var1, var2 = key[0], key[1]
                    controls = key[2:]
                    order = len(controls)

                    self.partial_corr_text.insert(
                        tk.END,
                        f"La corrélation partielle d'ordre {order} entre {var1} et {var2} "
                        f"(contrôlant pour {', '.join(controls)}) est {corr:.4f}.\n"
                        f"Interprétation: Relation linéaire entre {var1} et {var2} après avoir "
                        f"éliminé l'effet de {order} variable(s) de contrôle.\n"
                        f"Valeur absolue proche de 1 = relation forte indépendante des contrôles\n"
                        f"Valeur proche de 0 = relation expliquée par les variables de contrôle\n\n",
                    )

                self.partial_corr_text.insert(
                    tk.END,
                    "Les R² partiels montrent la contribution marginale de chaque variable "
                    "à l'explication de la variance de Y, après prise en compte des autres variables.\n"
                    "Un R² partiel élevé indique que la variable apporte une information "
                    "unique importante pour expliquer Y.\n",
                )

            messagebox.showinfo("Succès", "Calculs des corrélations et R² partiels terminés!")

        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur de calcul: {str(e)}")

    def klein_test(self):
        if self.model is None:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            self.klein_test_results = {}
            x_vars = [x.strip() for x in self.x_vars.get().split("+")]

            self.stat_tests_text.delete(1.0, tk.END)
            self.stat_tests_text.insert(
                tk.END, "=== TEST DE KLEIN POUR LA MULTICOLINÉARITÉ ===\n\n"
            )

            # Calcul du R² du modèle complet
            r_squared = self.model.rsquared

            # Calcul des corrélations simples entre variables explicatives
            corr_matrix = self.data[x_vars].corr()

            self.stat_tests_text.insert(tk.END, f"R² du modèle complet: {r_squared:.4f}\n\n")
            self.stat_tests_text.insert(
                tk.END, "Corrélations simples entre variables explicatives:\n"
            )
            self.stat_tests_text.insert(tk.END, str(corr_matrix) + "\n\n")

            # Vérification de la condition de Klein
            multicollinear_detected = False
            for i in range(len(x_vars)):
                for j in range(i + 1, len(x_vars)):
                    var1 = x_vars[i]
                    var2 = x_vars[j]
                    corr = corr_matrix.loc[var1, var2]
                    corr_squared = corr**2

                    self.klein_test_results[(var1, var2)] = (r_squared, corr_squared)

                    if r_squared < corr_squared:
                        multicollinear_detected = True
                        self.stat_tests_text.insert(
                            tk.END,
                            f"Problème potentiel entre {var1} et {var2}: R²={r_squared:.4f} < r²={corr_squared:.4f}\n",
                        )

            if not multicollinear_detected:
                self.stat_tests_text.insert(
                    tk.END, "Aucun problème de multicolinéarité détecté par le test de Klein\n"
                )

            # Hypothèses et règle de décision
            if self.show_hypothesis:
                self.stat_tests_text.insert(tk.END, "\n=== HYPOTHÈSES ET RÈGLE DE DÉCISION ===\n")
                self.stat_tests_text.insert(
                    tk.END,
                    "Test de Klein:\n"
                    "H0: Pas de multicolinéarité problématique entre les variables\n"
                    "H1: Présence de multicolinéarité problématique\n\n"
                    "Règle de décision:\n"
                    "Si R² du modèle < r² entre deux variables explicatives,\n"
                    "alors on rejette H0 et on conclut à une multicolinéarité problématique\n",
                )

            # Interprétation
            if self.show_interpretations:
                self.stat_tests_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                if multicollinear_detected:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le test de Klein a détecté une possible multicolinéarité problématique entre certaines variables.\n"
                        "Cela signifie que certaines variables explicatives sont fortement corrélées entre elles,\n"
                        "ce qui peut rendre les estimations des coefficients instables et difficiles à interpréter.\n"
                        "Solutions possibles:\n"
                        "- Supprimer une des variables corrélées\n"
                        "- Utiliser une analyse en composantes principales (ACP)\n"
                        "- Ajouter plus de données\n"
                        "- Utiliser la régression ridge\n",
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le test de Klein n'a pas détecté de multicolinéarité problématique.\n"
                        "Les variables explicatives ne semblent pas trop corrélées entre elles.\n",
                    )

            messagebox.showinfo("Succès", "Test de Klein effectué!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du test de Klein: {str(e)}")

    def farrar_glauber_test(self):
        if self.model is None:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            self.fg_test_results = {}
            x_vars = [x.strip() for x in self.x_vars.get().split("+")]

            self.stat_tests_text.delete(1.0, tk.END)
            self.stat_tests_text.insert(
                tk.END, "=== TEST DE FARRAR ET GLAUBER POUR LA MULTICOLINÉARITÉ ===\n\n"
            )

            # 1. Test du Chi² global
            X = self.data[x_vars]
            corr_matrix = X.corr()
            n = len(X)
            p = len(x_vars)

            det = np.linalg.det(corr_matrix)
            chi2 = -((n - 1) - (2 * p + 5) / 6) * np.log(det)
            df = p * (p - 1) // 2
            p_value = 1 - stats.chi2.cdf(chi2, df)

            self.fg_test_results["global"] = (chi2, df, p_value)

            self.stat_tests_text.insert(tk.END, "1. Test global du Chi²:\n")
            self.stat_tests_text.insert(
                tk.END, f"Déterminant de la matrice de corrélation: {det:.4f}\n"
            )
            self.stat_tests_text.insert(tk.END, f"Statistique Chi²: {chi2:.4f}\n")
            self.stat_tests_text.insert(tk.END, f"Degrés de liberté: {df}\n")
            self.stat_tests_text.insert(tk.END, f"p-value: {p_value:.4f}\n")

            if p_value < 0.05:
                self.stat_tests_text.insert(
                    tk.END, "Conclusion: Rejet de H0 - Multicolinéarité présente\n\n"
                )
            else:
                self.stat_tests_text.insert(
                    tk.END, "Conclusion: Non rejet de H0 - Pas de multicolinéarité globale\n\n"
                )

            # 2. Test F sur les régressions auxiliaires
            self.stat_tests_text.insert(tk.END, "2. Tests F sur les régressions auxiliaires:\n")

            f_tests = {}
            for i, var in enumerate(x_vars):
                other_vars = [v for v in x_vars if v != var]
                model = sm.OLS(self.data[var], sm.add_constant(self.data[other_vars])).fit()
                f_value = model.fvalue
                f_pvalue = model.f_pvalue
                f_tests[var] = (f_value, f_pvalue)

                self.stat_tests_text.insert(
                    tk.END,
                    f"Régression de {var} sur les autres variables:\n"
                    f"F-statistique: {f_value:.4f}, p-value: {f_pvalue:.4f}\n",
                )

                if f_pvalue < 0.05:
                    self.stat_tests_text.insert(
                        tk.END, f"Conclusion: {var} est liée aux autres variables\n\n"
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END, f"Conclusion: {var} n'est pas liée aux autres variables\n\n"
                    )

            self.fg_test_results["f_tests"] = f_tests

            # 3. Tests t sur les corrélations partielles
            self.stat_tests_text.insert(tk.END, "3. Tests t sur les corrélations partielles:\n")

            t_tests = {}
            for i in range(len(x_vars)):
                for j in range(i + 1, len(x_vars)):
                    var1 = x_vars[i]
                    var2 = x_vars[j]
                    other_vars = [v for v in x_vars if v not in [var1, var2]]

                    if len(other_vars) > 0:
                        res1 = sm.OLS(self.data[var1], sm.add_constant(self.data[other_vars])).fit()
                        res2 = sm.OLS(self.data[var2], sm.add_constant(self.data[other_vars])).fit()

                        residuals1 = res1.resid
                        residuals2 = res2.resid

                        corr, p_value = stats.pearsonr(residuals1, residuals2)
                        n = len(residuals1)
                        t_value = corr * np.sqrt((n - 2) / (1 - corr**2))

                        t_tests[(var1, var2)] = (t_value, p_value)

                        self.stat_tests_text.insert(
                            tk.END,
                            f"Corrélation partielle entre {var1} et {var2} | Contrôles: {', '.join(other_vars)}\n"
                            f"t-statistique: {t_value:.4f}, p-value: {p_value:.4f}\n",
                        )

                        if p_value < 0.05:
                            self.stat_tests_text.insert(
                                tk.END,
                                f"Conclusion: Corrélation significative entre {var1} et {var2}\n\n",
                            )
                        else:
                            self.stat_tests_text.insert(
                                tk.END,
                                f"Conclusion: Pas de corrélation significative entre {var1} et {var2}\n\n",
                            )

            self.fg_test_results["t_tests"] = t_tests

            # Hypothèses et règle de décision
            if self.show_hypothesis:
                self.stat_tests_text.insert(tk.END, "\n=== HYPOTHÈSES ET RÈGLES DE DÉCISION ===\n")
                self.stat_tests_text.insert(
                    tk.END,
                    "Test de Farrar-Glauber en trois étapes:\n\n"
                    "1. Test global du Chi²:\n"
                    "   H0: Matrice de corrélation = Matrice identité (pas de multicolinéarité)\n"
                    "   H1: Matrice de corrélation ≠ Matrice identité (multicolinéarité)\n"
                    "   Règle: Rejet de H0 si p-value < 0.05\n\n"
                    "2. Tests F sur les régressions auxiliaires:\n"
                    "   Pour chaque variable, on régresse sur les autres\n"
                    "   H0: Tous les coefficients = 0 (pas de relation linéaire)\n"
                    "   H1: Au moins un coefficient ≠ 0\n"
                    "   Règle: Rejet de H0 si p-value < 0.05\n\n"
                    "3. Tests t sur les corrélations partielles:\n"
                    "   H0: ρ(Xi,Xj|autres) = 0\n"
                    "   H1: ρ(Xi,Xj|autres) ≠ 0\n"
                    "   Règle: Rejet de H0 si p-value < 0.05\n",
                )

            # Interprétation
            if self.show_interpretations:
                self.stat_tests_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                if p_value < 0.05:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le test global indique la présence de multicolinéarité dans le modèle.\n"
                        "Les tests F et t permettent d'identifier quelles variables sont concernées.\n"
                        "Une multicolinéarité forte peut:\n"
                        "- Rendre les estimations des coefficients instables\n"
                        "- Augmenter les variances des estimateurs\n"
                        "- Rendre difficile l'interprétation des coefficients\n\n"
                        "Solutions possibles:\n"
                        "- Supprimer une des variables corrélées\n"
                        "- Utiliser des techniques de régularisation (ridge, lasso)\n"
                        "- Effectuer une analyse en composantes principales\n"
                        "- Collecter plus de données\n",
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le test global n'indique pas de problème de multicolinéarité.\n"
                        "Les variables explicatives semblent suffisamment indépendantes.\n",
                    )

            messagebox.showinfo("Succès", "Test de Farrar-Glauber effectué!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du test de Farrar-Glauber: {str(e)}")

    def determine_optimal_mix(self):
        if self.model is None:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            self.stat_tests_text.delete(1.0, tk.END)
            self.stat_tests_text.insert(tk.END, "=== DÉTERMINATION DU MIX OPTIMAL ===\n\n")

            x_vars = [x.strip() for x in self.x_vars.get().split("+")]

            # Calcul des élasticités
            y_mean = np.mean(self.y)
            x_means = [np.mean(self.data[var]) for var in x_vars]
            coefficients = [self.model.params[var] for var in x_vars]

            elasticities = []
            for coef, x_mean, var in zip(coefficients, x_means, x_vars):
                elasticity = coef * (x_mean / y_mean)
                elasticities.append(elasticity)
                self.stat_tests_text.insert(tk.END, f"Élasticité de {var}: {elasticity:.4f}\n")

            # Calcul des parts optimales
            sum_elasticities = sum(abs(e) for e in elasticities)
            optimal_shares = [abs(e) / sum_elasticities for e in elasticities]

            self.stat_tests_text.insert(tk.END, "\nParts optimales du mix:\n")
            for var, share in zip(x_vars, optimal_shares):
                self.stat_tests_text.insert(tk.END, f"{var}: {share * 100:.2f}%\n")

            # Formules
            if self.show_formulas:
                self.stat_tests_text.insert(tk.END, "\n=== FORMULES ===\n")
                self.stat_tests_text.insert(
                    tk.END,
                    "1. Calcul des élasticités:\n"
                    "   Elasticité = (∂Y/Y) / (∂X/X) = β * (X̄ / Ȳ)\n\n"
                    "2. Calcul des parts optimales:\n"
                    "   Part de Xi = |Elasticité Xi| / Σ|Elasticité Xj|\n",
                )

            # Interprétation
            if self.show_interpretations:
                self.stat_tests_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                self.stat_tests_text.insert(
                    tk.END,
                    "Le mix optimal représente la répartition idéale des ressources entre les différentes variables\n"
                    "explicatives pour maximiser l'impact sur la variable dépendante.\n"
                    "Plus l'élasticité d'une variable est élevée, plus son poids dans le mix optimal est important.\n"
                    "Ce calcul suppose que les variables sont indépendantes et que leurs effets sont additifs.\n"
                    "En pratique, d'autres contraintes (budgétaires, techniques) peuvent influencer le mix réel.\n",
                )

            messagebox.showinfo("Succès", "Mix optimal déterminé!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du calcul du mix optimal: {str(e)}")

    def test_variable_significance(self):
        if self.model is None:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            variable = simpledialog.askstring(
                "Test de significativité",
                "Entrez le nom de la variable à tester:",
                parent=self.root,
            )

            if variable and variable in self.model.params.index:
                self.stat_tests_text.delete(1.0, tk.END)
                self.stat_tests_text.insert(
                    tk.END, f"=== TEST DE SIGNIFICATIVITÉ DE {variable} ===\n\n"
                )

                # Récupérer les résultats du modèle
                idx = list(self.model.params.index).index(variable)
                coef = self.model.params[idx]
                std_err = self.model.bse[idx]
                t_value = self.model.tvalues[idx]
                p_value = self.model.pvalues[idx]

                # Afficher a ou beta selon la configuration
                coef_symbol = "a" if self.use_alpha else "β"

                self.stat_tests_text.insert(
                    tk.END, f"Coefficient {coef_symbol}_{variable}: {coef:.4f}\n"
                )
                self.stat_tests_text.insert(tk.END, f"Erreur standard: {std_err:.4f}\n")
                self.stat_tests_text.insert(tk.END, f"Statistique t: {t_value:.4f}\n")
                self.stat_tests_text.insert(tk.END, f"p-value: {p_value:.4f}\n")

                if p_value < self.alpha:
                    self.stat_tests_text.insert(
                        tk.END, f"\nConclusion: La variable est significative (p < {self.alpha})\n"
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END,
                        f"\nConclusion: La variable n'est pas significative (p ≥ {self.alpha})\n",
                    )

                # Hypothèses et règle de décision
                if self.show_hypothesis:
                    self.stat_tests_text.insert(
                        tk.END, "\n=== HYPOTHÈSES ET RÈGLE DE DÉCISION ===\n"
                    )
                    self.stat_tests_text.insert(
                        tk.END,
                        f"H0: Le coefficient {coef_symbol}_{variable} est égal à 0\n"
                        f"H1: Le coefficient {coef_symbol}_{variable} est différent de 0\n\n"
                        f"Règle de décision:\n"
                        f"Rejet de H0 si p-value < {self.alpha}\n"
                        f"On conclut alors que la variable a un effet significatif\n",
                    )

                # Valeur critique
                df = self.model.df_resid
                t_critical = stats.t.ppf(1 - self.alpha / 2, df)

                self.stat_tests_text.insert(
                    tk.END,
                    f"\nValeur critique t({df}) au seuil {self.alpha * 100}%: ±{t_critical:.4f}\n",
                )

                if abs(t_value) > t_critical:
                    self.stat_tests_text.insert(tk.END, "|t observé| > t critique ⇒ Rejet de H0\n")
                else:
                    self.stat_tests_text.insert(
                        tk.END, "|t observé| ≤ t critique ⇒ Non rejet de H0\n"
                    )

                # Intervalle de confiance
                ci_low = coef - t_critical * std_err
                ci_high = coef + t_critical * std_err

                self.stat_tests_text.insert(
                    tk.END,
                    f"\nIntervalle de confiance {100 * (1 - self.alpha)}%: [{ci_low:.4f}, {ci_high:.4f}]\n",
                )

                if ci_low <= 0 <= ci_high:
                    self.stat_tests_text.insert(
                        tk.END, "L'intervalle contient 0 ⇒ Non rejet de H0\n"
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END, "L'intervalle ne contient pas 0 ⇒ Rejet de H0\n"
                    )

                # Interprétation
                if self.show_interpretations:
                    self.stat_tests_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                    if p_value < self.alpha:
                        self.stat_tests_text.insert(
                            tk.END,
                            f"La variable {variable} a un effet significatif sur la variable dépendante.\n"
                            f"Une augmentation d'une unité de {variable} est associée à une variation de {coef:.4f} unités\n"
                            f"de la variable dépendante, toutes choses égales par ailleurs.\n"
                            f"Cette relation est statistiquement significative au seuil de {self.alpha * 100}%.\n",
                        )
                    else:
                        self.stat_tests_text.insert(
                            tk.END,
                            f"La variable {variable} ne semble pas avoir d'effet significatif sur la variable dépendante.\n"
                            f"L'estimation du coefficient n'est pas significativement différente de zéro.\n"
                            f"Cela peut signifier que:\n"
                            f"- La variable n'a vraiment pas d'effet\n"
                            f"- L'échantillon est trop petit pour détecter un effet\n"
                            f"- La spécification du modèle n'est pas appropriée\n",
                        )

                messagebox.showinfo("Succès", f"Test de significativité pour {variable} effectué!")
            else:
                messagebox.showwarning("Avertissement", "Variable non valide!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du test: {str(e)}")

    def test_global_significance(self):
        if self.model is None:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            self.stat_tests_text.delete(1.0, tk.END)
            self.stat_tests_text.insert(tk.END, "=== TEST DE SIGNIFICATIVITÉ GLOBALE ===\n\n")

            # Test F global
            f_value = self.model.fvalue
            f_pvalue = self.model.f_pvalue
            df_model = self.model.df_model
            df_resid = self.model.df_resid

            self.stat_tests_text.insert(tk.END, f"F-statistique: {f_value:.4f}\n")
            self.stat_tests_text.insert(tk.END, f"p-value: {f_pvalue:.4f}\n")
            self.stat_tests_text.insert(
                tk.END, f"Degrés de liberté: {df_model} (numérateur), {df_resid} (dénominateur)\n"
            )

            if f_pvalue < self.alpha:
                self.stat_tests_text.insert(
                    tk.END,
                    f"\nConclusion: Le modèle est globalement significatif (p < {self.alpha})\n",
                )
            else:
                self.stat_tests_text.insert(
                    tk.END,
                    f"\nConclusion: Le modèle n'est pas globalement significatif (p ≥ {self.alpha})\n",
                )

            # Hypothèses et règle de décision
            if self.show_hypothesis:
                self.stat_tests_text.insert(tk.END, "\n=== HYPOTHÈSES ET RÈGLE DE DÉCISION ===\n")
                self.stat_tests_text.insert(
                    tk.END,
                    "H0: Tous les coefficients (sauf la constante) sont égaux à 0\n"
                    "H1: Au moins un coefficient est différent de 0\n\n"
                    f"Règle de décision:\n"
                    f"Rejet de H0 si p-value < {self.alpha}\n"
                    "On conclut alors que le modèle est globalement significatif\n",
                )

            # Valeur critique
            f_critical = stats.f.ppf(1 - self.alpha, df_model, df_resid)

            self.stat_tests_text.insert(
                tk.END,
                f"\nValeur critique F({df_model},{df_resid}) au seuil {self.alpha * 100}%: {f_critical:.4f}\n",
            )

            if f_value > f_critical:
                self.stat_tests_text.insert(tk.END, "F observé > F critique ⇒ Rejet de H0\n")
            else:
                self.stat_tests_text.insert(tk.END, "F observé ≤ F critique ⇒ Non rejet de H0\n")

            # Analyse de variance
            # self.stat_tests_text.insert(tk.END, "\n=== ANALYSE DE VARIANCE (ANOVA) ===\n")

            # anova_table = anova_lm(self.model)
            # self.stat_tests_text.insert(tk.END, anova_table.to_string() + "\n")

            # Interprétation
            if self.show_interpretations:
                self.stat_tests_text.insert(tk.END, "\n=== INTERPRÉTATION ===\n")
                if f_pvalue < self.alpha:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le modèle dans son ensemble est statistiquement significatif.\n"
                        "Au moins une des variables explicatives a un effet significatif sur la variable dépendante.\n"
                        f"Cela valide globalement la pertinence du modèle au seuil de {self.alpha * 100}%.\n",
                    )
                else:
                    self.stat_tests_text.insert(
                        tk.END,
                        "Le modèle dans son ensemble n'est pas statistiquement significatif.\n"
                        f"Aucune des variables explicatives n'a d'effet significatif au seuil de {self.alpha * 100}%.\n"
                        "Le modèle pourrait ne pas être pertinent ou manquer de variables importantes.\n",
                    )

            messagebox.showinfo("Succès", "Test de significativité globale effectué!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du test: {str(e)}")

    def format_number(self, num):
        if isinstance(num, (int, float)):
            # Formater en notation standard plutôt qu'exponentielle
            return "{0:.{1}f}".format(num, self.decimal_places)
        return str(num)

    def import_data(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel", "*.xlsx *.xls"), ("CSV", "*.csv")]
        )
        if not file_path:
            return

        try:
            if file_path.endswith(".csv"):
                # Gestion CSV simplifiée
                self.data = pd.read_csv(file_path)
            else:
                # Configuration de la fenêtre d'importation
                import_dialog = tk.Toplevel(self.root)
                import_dialog.title("Importation Excel - Style EViews")
                import_dialog.geometry("900x600")

                # Variables pour les paramètres
                settings = {
                    "sheet": tk.StringVar(),
                    "header_row": tk.IntVar(value=0),
                    "data_start": tk.IntVar(value=0),
                    "col_names": tk.BooleanVar(value=True),
                    "index_col": tk.BooleanVar(value=False),
                    "na_values": tk.StringVar(value=""),
                }

                # Frame pour les contrôles
                control_frame = ttk.Frame(import_dialog)
                control_frame.pack(side="left", fill="y", padx=5, pady=5)

                # Frame pour l'aperçu
                preview_frame = ttk.Frame(import_dialog)
                preview_frame.pack(side="right", expand=True, fill="both", padx=5, pady=5)

                # Liste des feuilles
                xls = pd.ExcelFile(file_path)
                ttk.Label(control_frame, text="Feuille:").grid(row=0, column=0, sticky="e")
                sheet_cb = ttk.Combobox(
                    control_frame, textvariable=settings["sheet"], values=xls.sheet_names
                )
                sheet_cb.grid(row=0, column=1, sticky="ew", pady=2)
                sheet_cb.set(xls.sheet_names[0])

                # Options de lecture
                ttk.Label(control_frame, text="Ligne d'en-tête:").grid(row=1, column=0, sticky="e")
                header_spin = ttk.Spinbox(
                    control_frame, from_=0, to=100, textvariable=settings["header_row"]
                )
                header_spin.grid(row=1, column=1, sticky="ew", pady=2)

                ttk.Label(control_frame, text="Début des données:").grid(
                    row=2, column=0, sticky="e"
                )
                data_spin = ttk.Spinbox(
                    control_frame, from_=0, to=100, textvariable=settings["data_start"]
                )
                data_spin.grid(row=2, column=1, sticky="ew", pady=2)

                ttk.Checkbutton(
                    control_frame, text="Noms de colonnes", variable=settings["col_names"]
                ).grid(row=3, column=0, columnspan=2, sticky="w")
                ttk.Checkbutton(
                    control_frame, text="Index à partir colonne 1", variable=settings["index_col"]
                ).grid(row=4, column=0, columnspan=2, sticky="w")

                ttk.Label(control_frame, text="Valeurs manquantes:").grid(
                    row=5, column=0, sticky="e"
                )
                ttk.Entry(control_frame, textvariable=settings["na_values"]).grid(
                    row=5, column=1, sticky="ew", pady=2
                )

                # Treeview pour l'aperçu
                preview_tree = ttk.Treeview(preview_frame)
                vsb = ttk.Scrollbar(preview_frame, orient="vertical", command=preview_tree.yview)
                hsb = ttk.Scrollbar(preview_frame, orient="horizontal", command=preview_tree.xview)
                preview_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

                preview_tree.grid(row=0, column=0, sticky="nsew")
                vsb.grid(row=0, column=1, sticky="ns")
                hsb.grid(row=1, column=0, sticky="ew")

                preview_frame.grid_rowconfigure(0, weight=1)
                preview_frame.grid_columnconfigure(0, weight=1)

                # Fonction de mise à jour de l'aperçu
                def update_preview(event=None):
                    try:
                        read_params = {
                            "sheet_name": settings["sheet"].get(),
                            "header": settings["header_row"].get()
                            if settings["col_names"].get()
                            else None,
                            "skiprows": range(settings["data_start"].get()),
                            "index_col": 0 if settings["index_col"].get() else None,
                            "na_values": settings["na_values"].get().split(",")
                            if settings["na_values"].get()
                            else None,
                        }

                        preview_df = pd.read_excel(file_path, **read_params).head(20)

                        # Clear existing tree
                        preview_tree.delete(*preview_tree.get_children())
                        preview_tree["columns"] = list(preview_df.columns)

                        # Configure columns
                        for col in preview_df.columns:
                            preview_tree.heading(col, text=col)
                            preview_tree.column(col, width=100, anchor="center")

                        # Insert data
                        for _i, row in preview_df.iterrows():
                            preview_tree.insert("", "end", values=list(row))

                    except Exception as e:
                        preview_tree.delete(*preview_tree.get_children())
                        preview_tree.insert("", "end", values=[f"Erreur: {str(e)}"])

                # Lier les événements de modification
                for var in settings.values():
                    if isinstance(var, (tk.StringVar, tk.IntVar)):
                        var.trace_add("write", lambda *args: update_preview())

                settings["col_names"].trace_add("write", lambda *args: update_preview())
                settings["index_col"].trace_add("write", lambda *args: update_preview())

                # Boutons
                btn_frame = ttk.Frame(control_frame)
                btn_frame.grid(row=6, column=0, columnspan=2, pady=10)

                ttk.Button(btn_frame, text="Importer", command=import_dialog.destroy).pack(
                    side="left", padx=5
                )
                ttk.Button(btn_frame, text="Annuler", command=import_dialog.destroy).pack(
                    side="right", padx=5
                )

                # Initialiser l'aperçu
                update_preview()

                # Attendre la fermeture de la fenêtre
                self.root.wait_window(import_dialog)

                # Charger les données finales
                read_params = {
                    "sheet_name": settings["sheet"].get(),
                    "header": settings["header_row"].get() if settings["col_names"].get() else None,
                    "skiprows": range(settings["data_start"].get()),
                    "index_col": 0 if settings["index_col"].get() else None,
                    "na_values": settings["na_values"].get().split(",")
                    if settings["na_values"].get()
                    else None,
                }

                self.data = pd.read_excel(file_path, **read_params).dropna(how="all")

            # Mettre à jour l'interface
            self.update_interface_after_import()

        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur d'importation: {str(e)}")

    def update_interface_after_import(self):
        """Mettre à jour les menus après l'importation"""
        self.y_menu["values"] = list(self.data.columns)
        if len(self.data.columns) > 0:
            self.y_var.set(self.data.columns[0])
            if len(self.data.columns) > 1:
                self.x_vars.set(" + ".join(self.data.columns[1:]))

        self.display_data()
        messagebox.showinfo("Succès", f"{len(self.data)} observations importées!")

    def display_data(self):
        for i in self.data_tree.get_children():
            self.data_tree.delete(i)

        self.data_tree["columns"] = list(self.data.columns)
        self.data_tree["show"] = "headings"

        for col in self.data.columns:
            self.data_tree.heading(col, text=col)
            self.data_tree.column(col, width=100)

        for _index, row in self.data.iterrows():
            self.data_tree.insert("", "end", values=list(row))

    def clean_data(self):
        if self.data is None:
            messagebox.showwarning("Avertissement", "Importez d'abord des données!")
            return

        try:
            self.data = self.data.drop_duplicates()
            self.data = self.data.replace([np.inf, -np.inf], np.nan)

            for col in self.data.select_dtypes(include=[np.number]).columns:
                self.data[col].fillna(self.data[col].mean(), inplace=True)

            self.display_data()
            messagebox.showinfo("Succès", "Données nettoyées avec succès!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur de nettoyage: {str(e)}")

    def show_stats(self):
        if self.data is None:
            messagebox.showwarning("Avertissement", "Importez d'abord des données!")
            return

        stats = self.data.describe().to_string()
        messagebox.showinfo("Statistiques Descriptives", stats)

    def estimate_model(self):
        if self.data is None:
            messagebox.showwarning("Avertissement", "Importez d'abord des données!")
            return

        try:
            y = self.y_var.get()
            x_vars = [x.strip() for x in self.x_vars.get().split("+")]

            self.y = self.data[y]
            self.X = sm.add_constant(self.data[x_vars])

            self.model = sm.OLS(self.y, self.X).fit()

            self.show_results()
            self.update_prediction_entries()
            self.analyze_multicollinearity()
            self.calculate_partial_correlations()
            self.klein_test()
            self.farrar_glauber_test()

            messagebox.showinfo("Succès", "Modèle estimé avec succès!")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur d'estimation: {str(e)}")

    def show_results(self):
        self.results_text.delete(1.0, tk.END)

        # Résumé standard
        self.results_text.insert(tk.END, self.get_french_summary())

        # Formules mathématiques
        if self.show_formulas:
            self.results_text.insert(tk.END, "\n\n=== FORMULES MATHÉMATIQUES ===\n")
            self.results_text.insert(tk.END, self.get_model_formulas())

        # Matrices supplémentaires
        self.results_text.insert(tk.END, "\n\n=== MATRICES IMPORTANTES ===\n")

        # Matrice X
        self.results_text.insert(tk.END, "\nMatrice X:\n")
        self.results_text.insert(tk.END, str(self.X))

        # Matrice X'
        xt = self.X.T
        self.results_text.insert(tk.END, "\n\nMatrice X':\n")
        self.results_text.insert(tk.END, str(xt))

        # Matrice X'X
        xtx = np.dot(self.X.T, self.X)
        self.results_text.insert(tk.END, "\n\nMatrice X'X:\n")
        self.results_text.insert(tk.END, str(xtx))

        # Matrice X'Y
        xty = np.dot(self.X.T, self.y)
        self.results_text.insert(tk.END, "\n\nMatrice X'Y:\n")
        self.results_text.insert(tk.END, str(xty))

        # Matrice (X'X)^-1
        try:
            xtx_inv = np.linalg.inv(xtx)
            self.results_text.insert(tk.END, "\n\nMatrice (X'X)^-1:\n")
            self.results_text.insert(tk.END, str(xtx_inv))
        except:
            self.results_text.insert(tk.END, "\n\nMatrice (X'X) non inversible")

        # Variance des erreurs
        sigma2 = self.model.mse_resid
        self.results_text.insert(
            tk.END, f"\n\nVariance de l'erreur (σ²): {self.format_number(sigma2)}"
        )

        # Calcul de l'écart-type
        ecart_type = math.sqrt(sigma2)
        self.results_text.insert(
            tk.END, f"\n\nÉcart-type de l'erreur (σ): {self.format_number(ecart_type)}"
        )

        # Matrice variance-covariance
        self.results_text.insert(tk.END, "\n\nMatrice Variance-Covariance:\n")
        self.results_text.insert(tk.END, str(self.model.cov_params()))

        # Tableau complet des résultats
        self.results_text.insert(tk.END, "\n\n=== TABLEAU COMPLET DES RÉSULTATS ===\n")

        # Créer un DataFrame avec tous les résultats
        results_df = pd.DataFrame(
            {
                "Y_observé": self.y,
                "Y_prédit": self.model.fittedvalues,
                "Résidu": self.model.resid,
                "Résidu_carré": self.model.resid**2,
                "Résidu_studentisé": self.model.get_influence().resid_studentized_internal,
                "hi": self.model.get_influence().hat_matrix_diag,
            }
        )

        # Calculer les colonnes supplémentaires
        results_df["1-hi"] = 1 - results_df["hi"]
        # results_df['SCR'] = (results_df['Y_observé'] - results_df['Y_prédit'])**2
        results_df["SCR"] = (results_df["Y_prédit"] - np.mean(self.y)) ** 2
        results_df["SCT"] = (results_df["Y_observé"] - np.mean(self.y)) ** 2

        # Afficher le tableau
        self.results_text.insert(tk.END, results_df.to_string())

        # Tests de significativité
        self.results_text.insert(tk.END, "\n\n=== TESTS DE SIGNIFICATIVITÉ ===")

        # Test t pour chaque coefficient
        for i, var in enumerate(self.model.params.index):
            t_value = self.model.tvalues[i]
            p_value = self.model.pvalues[i]
            sig = (
                "***"
                if p_value < 0.01
                else "**"
                if p_value < 0.05
                else "*"
                if p_value < 0.1
                else ""
            )

            # Utiliser a ou beta selon la configuration
            coef_symbol = "a" if self.use_alpha else "β"
            var_name = f"{coef_symbol}_{var}" if var != "const" else "const"

            self.results_text.insert(
                tk.END,
                f"\n{var_name}: t = {self.format_number(t_value)}, p-value = {self.format_number(p_value)} {sig}",
            )

        # Test F global
        f_value = self.model.fvalue
        f_pvalue = self.model.f_pvalue
        self.results_text.insert(
            tk.END,
            f"\n\nTest F global: F = {self.format_number(f_value)}, p-value = {self.format_number(f_pvalue)}",
        )

        # Tests de diagnostic
        self.results_text.insert(tk.END, "\n\n=== TESTS DE DIAGNOSTIC ===")

        # Test de Durbin-Watson (autocorrélation)
        dw = durbin_watson(self.model.resid)
        self.results_text.insert(
            tk.END, f"\nDurbin-Watson (autocorrélation): {self.format_number(dw)}"
        )
        if dw < 1.5:
            self.results_text.insert(tk.END, " (Autocorrélation positive possible)")
        elif dw > 2.5:
            self.results_text.insert(tk.END, " (Autocorrélation négative possible)")
        else:
            self.results_text.insert(tk.END, " (Pas d'autocorrélation évidente)")

        # Test de Breusch-Pagan (hétéroscédasticité)
        bp_test = het_breuschpagan(self.model.resid, self.model.model.exog)
        self.results_text.insert(
            tk.END,
            f"\nBreusch-Pagan (hétéroscédasticité): LM = {self.format_number(bp_test[0])}, p-value = {self.format_number(bp_test[1])}",
        )

        # Test de normalité (Anderson-Darling)
        ad_test = normal_ad(self.model.resid)
        self.results_text.insert(
            tk.END,
            f"\nAnderson-Darling (normalité): A² = {self.format_number(ad_test[0])}, p-value = {self.format_number(ad_test[1])}",
        )

        # Test de Jarque-Bera (normalité)
        jb_test = jarque_bera(self.model.resid)
        self.results_text.insert(
            tk.END,
            f"\nJarque-Bera (normalité): JB = {self.format_number(jb_test[0])}, p-value = {self.format_number(jb_test[1])}",
        )

        # Interprétations
        if self.show_interpretations:
            self.results_text.insert(tk.END, "\n\n=== INTERPRÉTATIONS ===\n")
            self.results_text.insert(tk.END, self.get_interpretations())

        # Étapes de calcul
        if self.show_calculation_steps:
            self.results_text.insert(tk.END, "\n\n=== ÉTAPES DE CALCUL ===\n")
            self.results_text.insert(tk.END, self.get_calculation_steps())

    def get_calculation_steps(self):
        if not self.model:
            return ""

        # Utiliser a ou beta selon la configuration
        coef_symbol = "a" if self.use_alpha else "β"
        # Calcul de la matrice hat
        X = self.X.values if hasattr(self.X, "values") else self.X
        XtX_inv = np.linalg.inv(np.dot(X.T, X))
        H = np.dot(X, np.dot(XtX_inv, X.T))
        h_diag = np.diag(H)  # Éléments diagonaux (leviers)
        steps = (
            f"1. Estimation des coefficients ({coef_symbol}):\n"
            f"   {coef_symbol}̂ = (X'X)⁻¹X'y\n"
            f"   X'X = \n{np.dot(self.X.T, self.X)}\n\n"
            f"   X'y = \n{np.dot(self.X.T, self.y)}\n\n"
            f"   (X'X)⁻¹ = \n{np.linalg.inv(np.dot(self.X.T, self.X))}\n\n"
            f"   {coef_symbol}̂ = \n{self.model.params}\n\n"
            "2. Calcul des valeurs prédites (ŷ):\n"
            f"   ŷ = X{coef_symbol}̂\n"
            f"   Exemple pour la première observation: \n{self.X.iloc[0]} * {coef_symbol}̂ = {np.dot(self.X.iloc[0], self.model.params)}\n\n"
            "3. Calcul des résidus (e):\n"
            "   e = y - ŷ\n"
            f"   Exemple pour la première observation: {self.y.iloc[0]} - {self.model.fittedvalues[0]} = {self.model.resid[0]}\n\n"
            "4. Calcul de la variance des résidus (σ²):\n"
            "   σ̂² = e'e / (n - k)\n"
            f"   e'e = {np.dot(self.model.resid, self.model.resid)}\n"
            f"   n = {self.model.nobs}, k = {len(self.model.params)}\n"
            f"   σ̂² = {self.model.mse_resid}\n"
            f"   σ = {np.sqrt(self.model.mse_resid)}\n\n"  # Écart-type des résidus (racine carrée de la variance)
            f"5. Calcul des écarts-types des coefficients:\n"
            f"   Ω = Var({coef_symbol}̂) = σ̂²(X'X)⁻¹\n"
            f"   se({coef_symbol}̂_j) = √(Var({coef_symbol}̂)_jj)\n"
            f"   Matrice Var({coef_symbol}̂): \n{self.model.cov_params()}\n"
            f"   Écarts-types: \n{self.model.bse}\n\n"
            f"6. Calcul des statistiques t:\n"
            f"   t_j = {coef_symbol}̂_j / se({coef_symbol}̂_j)\n"
            f"   Statistiques t: \n{self.model.tvalues}\n\n"
            "7. Calcul du R²:\n"
            "   SCT = Σ(y_i - ȳ)²\n"
            "   SCE = Σ(y_i - ŷ_i)²\n"
            "   R² = 1 - SCE/SCT\n"
            f"   SCT = {np.sum((self.y - np.mean(self.y)) ** 2)}\n"
            f"   SCE = {np.sum(self.model.resid**2)}\n"
            f"   R² = {self.model.rsquared}\n\n"
            "8. Calcul du R² ajusté:\n"
            "   R²_adj = 1 - (1-R²)(n-1)/(n-k-1)\n"
            f"   R²_adj = {self.model.rsquared_adj}\n\n"
            "9. Calcul de la F-statistique:\n"
            "   F = [(SCT - SCE)/k] / [SCE/(n - k - 1)]\n"
            f"   F = {self.model.fvalue}\n\n"
            "10. Calcul de la matrice de projection (hat matrix H):\n"
            "   H = X(X'X)⁻¹X'\n"
            f"   Dimensions de H: {H.shape}\n"
            "   Exemple de sous-matrice de H (5 premières lignes/colonnes):\n"
            f"{H[:5, :5]}\n"
            "   Éléments diagonaux (leviers) h_ii (5 premiers):\n"
            f"{h_diag[:5]}\n"
            "   Trace de H (somme des leviers): tr(H) = k = \n"
            f"{np.trace(H)} (nombre de variables explicatives + constante)\n\n"
        )

        return steps

    def get_french_summary(self):
        if not self.model:
            return "Aucun modèle estimé"

        now = datetime.now()
        date_str = now.strftime("%a, %d %b %Y")
        time_str = now.strftime("%H:%M:%S")

        # Utiliser a ou beta selon la configuration
        coef_symbol = "a" if self.use_alpha else "β"

        # Modifier le tableau des coefficients pour utiliser a au lieu de beta
        summary_tables = self.model.summary().tables
        if len(summary_tables) > 1:
            coef_table = summary_tables[1].as_text()
            if self.use_alpha:
                coef_table = coef_table.replace("beta", "alpha").replace("const", "a0")

        summary = (
            "Variable Dépendante:           {dep_var}   R-carré:                       {rsquared}\n"
            "Modèle:                       OLS   R-carré ajusté:              {rsquared_adj}\n"
            "Méthode:          Moindres Carrés   F-statistique:               {fvalue}\n"
            "Date:             {date}   Prob (F-statistique):          {f_pvalue}\n"
            "Heure:                 {time}   Log-vraisemblance:            {ll}\n"
            "Nb. Observations:      {nobs}   AIC:                          {aic}\n"
            "Df Résidus:          {df_resid}   BIC:                          {bic}\n"
            "Df Modèle:            {df_model}                                         \n"
            "Covariance Type:   nonrobuste                                         \n"
            "==============================================================================\n"
            "{params_table}\n"
            "==============================================================================\n"
        ).format(
            dep_var=self.model.model.endog_names,
            rsquared=self.format_number(self.model.rsquared),
            rsquared_adj=self.format_number(self.model.rsquared_adj),
            fvalue=self.format_number(self.model.fvalue),
            date=date_str,
            time=time_str,
            ll=self.format_number(self.model.llf),
            nobs=self.model.nobs,
            aic=self.format_number(self.model.aic),
            bic=self.format_number(self.model.bic),
            df_resid=self.model.df_resid,
            df_model=self.model.df_model,
            f_pvalue=self.format_number(self.model.f_pvalue),
            params_table=coef_table
            if "coef_table" in locals()
            else self.model.summary().tables[1].as_text(),
        )

        return summary

    def get_model_formulas(self):
        # Utiliser a ou beta selon la configuration
        coef_symbol = "a" if self.use_alpha else "β"

        formulas = (
            f"1. Modèle de régression:\n"
            f"   Y = {coef_symbol}₀ + {coef_symbol}₁X₁ + {coef_symbol}₂X₂ + ... + {coef_symbol}ₖXₖ + ε\n\n"
            f"2. Estimateur MCO:\n"
            f"   {coef_symbol}̂ = (X'X)⁻¹X'y\n\n"
            "3. Matrice de projection (hat matrix):\n"
            "   H = X(X'X)⁻¹X'\n"
            "   ŷ = Hy = Xβ̂ (valeurs prédites)\n"
            "   h_ii = éléments diagonaux de H (levier)\n\n"
            "4. Variance des résidus:\n"
            f"   σ̂² = (y - X{coef_symbol}̂)'(y - X{coef_symbol}̂) / (n - k) = ε̂'ε̂ / (n - k)\n\n"
            f"5. Matrice variance-covariance:\n"
            f"   Var({coef_symbol}̂) = σ̂²(X'X)⁻¹\n\n"
            f"6. Statistique t:\n"
            f"   t = {coef_symbol}̂ᵢ / se({coef_symbol}̂ᵢ)\n\n"
            "7. Statistique F:\n"
            "   F = [(SCE₀ - SCE₁)/q] / [SCE₁/(n - k)]\n"
            "   où SCE = somme des carrés des résidus\n\n"
            "8. R² et R² ajusté:\n"
            "   R² = 1 - SCE/SCT\n"
            "   R²_adj = 1 - (1-R²)(n-1)/(n-k-1)\n\n"
            "9. Corrélation partielle:\n"
            "   ρ(X1,X2|Z) = corr(e1, e2)\n"
            "   où e1 = résidus de X1~Z, e2 = résidus de X2~Z\n\n"
            "10. Test de Klein (multicolinéarité):\n"
            "    Si R²_y < r²_xi,xj ⇒ multicolinéarité\n\n"
            "11. Test de Farrar-Glauber (multicolinéarité):\n"
            "    a) Test χ² global\n"
            "    b) Tests F sur régressions auxiliaires\n"
            "    c) Tests t sur corrélations partielles\n"
        )

        return formulas

    def get_interpretations(self):
        if not self.model:
            return ""

        # Utiliser a ou beta selon la configuration
        coef_symbol = "a" if self.use_alpha else "β"

        interpretations = []

        # Interprétation R²
        rsq = self.model.rsquared
        if rsq > 0.9:
            interp = f"Le R² de {self.format_number(rsq)} indique que le modèle explique plus de 90% de la variabilité de Y, ce qui suggère un excellent ajustement."
        elif rsq > 0.7:
            interp = f"Le R² de {self.format_number(rsq)} indique que le modèle explique une grande partie de la variabilité de Y."
        elif rsq > 0.5:
            interp = f"Le R² de {self.format_number(rsq)} indique que le modèle explique environ la moitié de la variabilité de Y."
        else:
            interp = f"Le R² de {self.format_number(rsq)} est relativement faible, ce qui suggère que le modèle explique peu de la variabilité de Y."
        interpretations.append(interp)

        # Interprétation F-statistique
        f_pval = self.model.f_pvalue
        if f_pval < 0.05:
            interp = (
                f"La F-statistique significative (p-value = {self.format_number(f_pval)}) "
                "indique que le modèle dans son ensemble est statistiquement significatif."
            )
        else:
            interp = (
                f"La F-statistique non significative (p-value = {self.format_number(f_pval)}) "
                "suggère que le modèle n'est pas meilleur qu'un modèle sans variables explicatives."
            )
        interpretations.append(interp)

        # Interprétation des coefficients
        for i, var in enumerate(self.model.params.index):
            if var == "const":
                continue

            coef = self.model.params[i]
            pval = self.model.pvalues[i]

            var_name = f"{coef_symbol}_{var}"

            if pval < 0.05:
                interp = (
                    f"La variable {var_name} est significative (p-value = {self.format_number(pval)}). "
                    f"Une augmentation d'une unité de {var} est associée à une variation de {self.format_number(coef)} "
                    "unités de Y, toutes choses égales par ailleurs."
                )
            else:
                interp = (
                    f"La variable {var_name} n'est pas significative (p-value = {self.format_number(pval)}), "
                    "ce qui suggère qu'elle pourrait ne pas avoir d'impact sur Y dans ce modèle."
                )
            interpretations.append("\n" + interp)

        # Interprétation Durbin-Watson
        dw = durbin_watson(self.model.resid)
        if dw < 1.5:
            interp = "\nLa valeur du test de Durbin-Watson suggère une possible autocorrélation positive des résidus, ce qui peut indiquer une spécification incorrecte du modèle."
        elif dw > 2.5:
            interp = "\nLa valeur du test de Durbin-Watson suggère une possible autocorrélation négative des résidus."
        else:
            interp = "\nLe test de Durbin-Watson ne détecte pas d'autocorrélation significative des résidus."
        interpretations.append(interp)

        # Interprétation hétéroscédasticité
        bp_test = het_breuschpagan(self.model.resid, self.model.model.exog)
        if bp_test[1] < 0.05:
            interp = "\nLe test de Breusch-Pagan rejette l'hypothèse d'homoscédasticité (p-value < 0.05), ce qui suggère la présence d'hétéroscédasticité."
        else:
            interp = "\nLe test de Breusch-Pagan ne rejette pas l'hypothèse d'homoscédasticité, ce qui est une bonne nouvelle pour les propriétés des estimateurs MCO."
        interpretations.append(interp)

        # Interprétation normalité
        ad_test = normal_ad(self.model.resid)
        if ad_test[1] < 0.05:
            interp = "\nLe test d'Anderson-Darling rejette l'hypothèse de normalité des résidus (p-value < 0.05), ce qui peut affecter la validité des tests d'hypothèse."
        else:
            interp = "\nLe test d'Anderson-Darling ne rejette pas l'hypothèse de normalité des résidus, ce qui est favorable pour l'inférence statistique."
        interpretations.append(interp)

        # Interprétation Jarque-Bera
        jb_test = jarque_bera(self.model.resid)
        if jb_test[1] < 0.05:
            interp = "\nLe test de Jarque-Bera rejette l'hypothèse de normalité des résidus (p-value < 0.05), confirmant la non-normalité."
        else:
            interp = "\nLe test de Jarque-Bera ne rejette pas l'hypothèse de normalité des résidus."
        interpretations.append(interp)

        return "\n".join(interpretations)

    def analyze_multicollinearity(self):
        if self.model is None or len(self.model.params) < 2:
            self.multicoll_text.delete(1.0, tk.END)
            self.multicoll_text.insert(
                tk.END, "Pas assez de variables pour analyser la multicolinéarité"
            )
            return

        self.multicoll_text.delete(1.0, tk.END)

        # Facteurs d'inflation de la variance (VIF)
        self.multicoll_text.insert(tk.END, "=== FACTEURS D'INFLATION DE LA VARIANCE (VIF) ===\n")
        vif_data = []
        for i, col in enumerate(self.X.columns):
            if col == "const":
                continue
            vif = variance_inflation_factor(self.X.values, i)
            vif_data.append((col, vif))

        # Trier par VIF décroissant
        vif_data.sort(key=lambda x: x[1], reverse=True)

        for var, vif in vif_data:
            interpretation = ""
            if vif > 10:
                interpretation = " (Multicolinéarité forte)"
            elif vif > 5:
                interpretation = " (Multicolinéarité modérée)"

            self.multicoll_text.insert(
                tk.END, f"{var}: VIF = {self.format_number(vif)}{interpretation}\n"
            )

        # Indice de condition et valeurs propres
        self.multicoll_text.insert(tk.END, "\n=== INDICE DE CONDITION ET VALEURS PROPRES ===\n")
        X = self.X.drop("const", axis=1) if "const" in self.X.columns else self.X
        X_centered = X - X.mean()
        eigvals = np.linalg.eigvals(X_centered.T @ X_centered)
        condition_indices = np.sqrt(eigvals.max() / eigvals)

        self.multicoll_text.insert(
            tk.END, f"Valeurs propres: {[self.format_number(v) for v in eigvals]}\n"
        )
        self.multicoll_text.insert(
            tk.END,
            f"Indices de condition: {[self.format_number(ci) for ci in condition_indices]}\n",
        )

        # Interprétation
        max_ci = max(condition_indices)
        if max_ci > 30:
            self.multicoll_text.insert(
                tk.END,
                f"\nL'indice de condition maximal est {self.format_number(max_ci)}, "
                "ce qui indique une multicolinéarité forte.\n",
            )
        elif max_ci > 10:
            self.multicoll_text.insert(
                tk.END,
                f"\nL'indice de condition maximal est {self.format_number(max_ci)}, "
                "ce qui indique une multicolinéarité modérée.\n",
            )
        else:
            self.multicoll_text.insert(
                tk.END,
                f"\nL'indice de condition maximal est {self.format_number(max_ci)}, "
                "ce qui suggère une multicolinéarité faible.\n",
            )

        # Matrice de corrélation
        self.multicoll_text.insert(tk.END, "\n=== MATRICE DE CORRÉLATION ===\n")
        corr_matrix = X.corr()
        self.multicoll_text.insert(tk.END, str(corr_matrix))

        # Formules
        if self.show_formulas:
            self.multicoll_text.insert(tk.END, "\n\n=== FORMULES MULTICOLINÉARITÉ ===\n")
            self.multicoll_text.insert(
                tk.END,
                "1. Facteur d'Inflation de la Variance (VIF):\n"
                "   VIF_j = 1 / (1 - R²_j)\n"
                "   où R²_j est le R² de la régression de X_j sur les autres variables explicatives\n\n"
                "2. Indice de Condition:\n"
                "   κ_j = √(λ_max / λ_j)\n"
                "   où λ sont les valeurs propres de X'X\n",
            )

        # Interprétation
        if self.show_interpretations:
            self.multicoll_text.insert(tk.END, "\n\n=== INTERPRÉTATION ===\n")
            if any(vif > 10 for var, vif in vif_data) or max_ci > 30:
                self.multicoll_text.insert(
                    tk.END,
                    "La multicolinéarité semble forte dans le modèle.\n"
                    "Cela peut entraîner:\n"
                    "- Des estimations de coefficients instables\n"
                    "- Des erreurs standard élevées\n"
                    "- Des difficultés à évaluer l'importance relative des variables\n\n"
                    "Solutions possibles:\n"
                    "- Supprimer certaines variables corrélées\n"
                    "- Utiliser la régression ridge ou lasso\n"
                    "- Effectuer une analyse en composantes principales\n"
                    "- Collecter plus de données\n",
                )
            elif any(vif > 5 for var, vif in vif_data) or max_ci > 10:
                self.multicoll_text.insert(
                    tk.END,
                    "Une certaine multicolinéarité est présente mais pas alarmante.\n"
                    "Les résultats peuvent être interprétés avec prudence.\n",
                )
            else:
                self.multicoll_text.insert(
                    tk.END,
                    "La multicolinéarité semble faible.\n"
                    "Les estimations des coefficients sont fiables.\n",
                )

        # Étapes de calcul si demandé
        if self.show_calculation_steps:
            self.multicoll_text.insert(tk.END, "\n\n=== ÉTAPES DE CALCUL VIF ===\n")
            for var, vif in vif_data:
                if var == "const":
                    continue

                other_vars = [v for v in self.X.columns if v not in [var, "const"]]
                if len(other_vars) > 0:
                    model = sm.OLS(self.X[var], sm.add_constant(self.X[other_vars])).fit()
                    rsq = model.rsquared
                    self.multicoll_text.insert(
                        tk.END,
                        f"Pour {var}, R² de la régression sur {', '.join(other_vars)} = {rsq:.4f}\n"
                        f"VIF = 1 / (1 - {rsq:.4f}) = {1 / (1 - rsq):.4f}\n\n",
                    )

    def update_prediction_entries(self):
        # Supprimer les anciens widgets
        for widget in self.pred_new_frame.winfo_children():
            widget.destroy()

        if self.model:
            x_vars = [x.strip() for x in self.x_vars.get().split("+")]
            self.pred_entries = {}

            # Ajouter les champs pour chaque variable
            for i, var in enumerate(x_vars):
                ttk.Label(self.pred_new_frame, text=f"{var}:").grid(row=i, column=0, padx=5, pady=2)
                entry = ttk.Entry(self.pred_new_frame)
                entry.grid(row=i, column=1, padx=5, pady=2)
                self.pred_entries[var] = entry

            # Bouton de calcul
            ttk.Button(
                self.pred_new_frame,
                text="Calculer Nouvelle Prédiction",
                command=self.calculate_new_prediction,
            ).grid(row=len(x_vars), column=0, columnspan=2, pady=5)

            # Champ pour afficher le résultat
            self.new_pred_result = ttk.Label(self.pred_new_frame, text="")
            self.new_pred_result.grid(row=len(x_vars) + 1, column=0, columnspan=2)

    def calculate_new_prediction(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            values = {"const": 1}
            for var in self.pred_entries:
                values[var] = float(self.pred_entries[var].get())

            pred_df = pd.DataFrame([values], columns=self.X.columns)

            prediction = self.model.get_prediction(pred_df)
            pred_mean = prediction.predicted_mean[0]
            pred_ci = prediction.conf_int(alpha=self.alpha)[0]  # Utilisation de self.alpha
            pred_std = prediction.se_mean[0]

            result = (
                f"Prédiction: y = {self.format_number(pred_mean)}\n"
                f"Écart-type de prévision: {self.format_number(pred_std)}\n"
                f"Intervalle de confiance {int(100 * (1 - self.alpha))}%: "  # Affichage du pourcentage correct
                f"[{self.format_number(pred_ci[0])}, {self.format_number(pred_ci[1])}]"
            )
            self.new_pred_result.config(text=result)

        except ValueError:
            messagebox.showerror("Erreur", "Veuillez entrer des valeurs numériques valides!")

    def calculate_prediction(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        try:
            # Utiliser les dernières valeurs des données
            last_row = self.X.iloc[-1].copy()
            last_row["const"] = 1  # Assurer que la constante est incluse

            pred_df = pd.DataFrame([last_row], columns=self.X.columns)

            prediction = self.model.get_prediction(pred_df)
            pred_mean = prediction.predicted_mean[0]
            pred_ci = prediction.conf_int(alpha=self.alpha)[0]  # Utilisation de self.alpha
            pred_std = prediction.se_mean[0]

            result = (
                f"Prédiction pour les dernières valeurs:\n"
                f"Y = {self.format_number(pred_mean)}\n"
                f"Écart-type: {self.format_number(pred_std)}\n"
                f"Intervalle de confiance {int(100 * (1 - self.alpha))}%: "  # Affichage du pourcentage correct
                f"[{self.format_number(pred_ci[0])}, {self.format_number(pred_ci[1])}]"
            )
            self.pred_result.config(text=result)

        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur de prédiction: {str(e)}")

    def show_diagnostic(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Estimez d'abord un modèle!")
            return

        diag_type = self.diag_type.get()
        self.figure.clf()
        ax = self.figure.add_subplot(111)

        if diag_type == "Résidus vs Ajustés":
            ax.scatter(self.model.fittedvalues, self.model.resid)
            ax.axhline(y=0, color="r", linestyle="--")
            ax.set_xlabel("Valeurs ajustées")
            ax.set_ylabel("Résidus")
            ax.set_title("Résidus vs Valeurs Ajustées")

        elif diag_type == "QQ Plot":
            sm.qqplot(self.model.resid, line="s", ax=ax)
            ax.set_title("QQ Plot des Résidus")

        elif diag_type == "Leverage":
            sm.graphics.influence_plot(self.model, ax=ax)
            ax.set_title("Leverage Plot")

        elif diag_type == "Histogramme Résidus":
            sns.histplot(self.model.resid, kde=True, ax=ax)
            ax.set_xlabel("Résidus")
            ax.set_title("Distribution des Résidus")

        elif diag_type == "Corrélation":
            corr = self.data.corr()
            sns.heatmap(corr, annot=True, fmt=f".{self.decimal_places}f", ax=ax)
            ax.set_title("Matrice de Corrélation")

        elif diag_type == "Autocorrélation":
            plot_acf(self.model.resid, ax=ax)
            ax.set_title("Fonction d'Autocorrélation")

        elif diag_type == "Autocorrélation Partielle":
            plot_pacf(self.model.resid, ax=ax)
            ax.set_title("Fonction d'Autocorrélation Partielle")

        self.canvas.draw()

    def export_to_excel(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Aucun modèle à exporter!")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx"), ("Tous fichiers", "*.*")]
        )

        if file_path:
            try:
                with pd.ExcelWriter(file_path) as writer:
                    # Données
                    if self.include_data.get():
                        self.data.to_excel(writer, sheet_name="Données")

                    # Résumé modèle
                    if self.include_summary.get():
                        summary_df = pd.DataFrame(
                            {
                                "Statistique": [
                                    "R²",
                                    "R² ajusté",
                                    "F-statistique",
                                    "Prob(F-statistique)",
                                    "Log-vraisemblance",
                                    "AIC",
                                    "BIC",
                                    "Nb. Observations",
                                    "Seuil alpha",
                                ],
                                "Valeur": [
                                    self.model.rsquared,
                                    self.model.rsquared_adj,
                                    self.model.fvalue,
                                    self.model.f_pvalue,
                                    self.model.llf,
                                    self.model.aic,
                                    self.model.bic,
                                    self.model.nobs,
                                    self.alpha,
                                ],
                            }
                        )
                        summary_df.to_excel(writer, sheet_name="Résumé", index=False)

                        # Coefficients
                        params_df = pd.DataFrame(
                            {
                                "Variable": self.model.params.index,
                                "Coefficient": self.model.params.values,
                                "Ecart-type": self.model.bse.values,
                                "t": self.model.tvalues.values,
                                "P>|t|": self.model.pvalues.values,
                                "Significatif": [
                                    "Oui" if p < self.alpha else "Non"
                                    for p in self.model.pvalues.values
                                ],
                            }
                        )
                        params_df.to_excel(writer, sheet_name="Coefficients", index=False)

                    # Matrices
                    if self.include_matrices.get():
                        X_df = pd.DataFrame(self.X)
                        X_df.to_excel(writer, sheet_name="Matrice_X")

                        xtx = np.dot(self.X.T, self.X)
                        pd.DataFrame(xtx, index=self.X.columns, columns=self.X.columns).to_excel(
                            writer, sheet_name="Matrice_XtX"
                        )

                        xty = np.dot(self.X.T, self.y)
                        pd.DataFrame(xty, index=self.X.columns, columns=["X'Y"]).to_excel(
                            writer, sheet_name="Matrice_XtY"
                        )

                        try:
                            xtx_inv = np.linalg.inv(xtx)
                            pd.DataFrame(
                                xtx_inv, index=self.X.columns, columns=self.X.columns
                            ).to_excel(writer, sheet_name="Matrice_XtX_inv")
                        except:
                            pass

                        pd.DataFrame(self.model.cov_params()).to_excel(
                            writer, sheet_name="Matrice_VarCov"
                        )

                    # Diagnostics
                    if self.include_diagnostics.get():
                        # Tests
                        tests_df = pd.DataFrame(
                            {
                                "Test": [
                                    "Durbin-Watson",
                                    "Breusch-Pagan",
                                    "Anderson-Darling",
                                    "Jarque-Bera",
                                ],
                                "Statistique": [
                                    durbin_watson(self.model.resid),
                                    het_breuschpagan(self.model.resid, self.model.model.exog)[0],
                                    normal_ad(self.model.resid)[0],
                                    jarque_bera(self.model.resid)[0],
                                ],
                                "p-value": [
                                    "",
                                    het_breuschpagan(self.model.resid, self.model.model.exog)[1],
                                    normal_ad(self.model.resid)[1],
                                    jarque_bera(self.model.resid)[1],
                                ],
                                "Seuil alpha": [self.alpha, self.alpha, self.alpha, self.alpha],
                                "Conclusion": [
                                    "",
                                    "Hétéroscédasticité"
                                    if het_breuschpagan(self.model.resid, self.model.model.exog)[1]
                                    < self.alpha
                                    else "OK",
                                    "Non-normalité"
                                    if normal_ad(self.model.resid)[1] < self.alpha
                                    else "OK",
                                    "Non-normalité"
                                    if jarque_bera(self.model.resid)[1] < self.alpha
                                    else "OK",
                                ],
                            }
                        )
                        tests_df.to_excel(writer, sheet_name="Tests_Diagnostic", index=False)

                        # Résidus
                        resid_df = pd.DataFrame(
                            {
                                "Y_observé": self.y,
                                "Y_prédit": self.model.fittedvalues,
                                "Résidu": self.model.resid,
                                "Résidu_standardisé": self.model.get_influence().resid_studentized_internal,
                                "Leverage": self.model.get_influence().hat_matrix_diag,
                            }
                        )
                        resid_df.to_excel(writer, sheet_name="Résidus", index=False)

                    # Multicolinéarité
                    vif_data = []
                    for i, col in enumerate(self.X.columns):
                        if col == "const":
                            continue
                        vif = variance_inflation_factor(self.X.values, i)
                        vif_data.append((col, vif))

                    pd.DataFrame(vif_data, columns=["Variable", "VIF"]).to_excel(
                        writer, sheet_name="Multicolinearite", index=False
                    )

                    # Corrélations partielles
                    if self.include_partial_corr.get() and self.partial_corr_results:
                        partial_corr_df = pd.DataFrame(
                            {
                                "Variables": [
                                    f"{k[0]} vs {k[1]}" for k in self.partial_corr_results
                                ],
                                "Corrélation_Partielle": list(self.partial_corr_results.values()),
                            }
                        )
                        partial_corr_df.to_excel(
                            writer, sheet_name="Corrélations_Partielles", index=False
                        )

                    # Tests de Klein
                    if self.include_stat_tests.get() and self.klein_test_results:
                        klein_df = pd.DataFrame(
                            {
                                "Variables": [f"{k[0]} vs {k[1]}" for k in self.klein_test_results],
                                "R²_modèle": [v[0] for v in self.klein_test_results.values()],
                                "r²_variables": [v[1] for v in self.klein_test_results.values()],
                                "Seuil alpha": [
                                    self.alpha for _ in self.klein_test_results.values()
                                ],
                                "Problème": [
                                    "Oui" if v[0] < v[1] else "Non"
                                    for v in self.klein_test_results.values()
                                ],
                            }
                        )
                        klein_df.to_excel(writer, sheet_name="Test_Klein", index=False)

                    # Tests de Farrar-Glauber
                    if self.include_stat_tests.get() and self.fg_test_results:
                        # Test global
                        fg_global_df = pd.DataFrame(
                            {
                                "Test": ["Chi² global"],
                                "Statistique": [self.fg_test_results["global"][0]],
                                "Degrés liberté": [self.fg_test_results["global"][1]],
                                "p-value": [self.fg_test_results["global"][2]],
                                "Seuil alpha": [self.alpha],
                                "Conclusion": [
                                    "Multicolinéarité"
                                    if self.fg_test_results["global"][2] < self.alpha
                                    else "Pas de multicolinéarité"
                                ],
                            }
                        )
                        fg_global_df.to_excel(writer, sheet_name="Test_FG_Global", index=False)

                        # Tests F
                        fg_f_df = pd.DataFrame(
                            {
                                "Variable": list(self.fg_test_results["f_tests"].keys()),
                                "F-statistique": [
                                    v[0] for v in self.fg_test_results["f_tests"].values()
                                ],
                                "p-value": [v[1] for v in self.fg_test_results["f_tests"].values()],
                                "Seuil alpha": [
                                    self.alpha for _ in self.fg_test_results["f_tests"].values()
                                ],
                                "Conclusion": [
                                    "Liée aux autres" if v[1] < self.alpha else "Non liée"
                                    for v in self.fg_test_results["f_tests"].values()
                                ],
                            }
                        )
                        fg_f_df.to_excel(writer, sheet_name="Test_FG_F", index=False)

                        # Tests t
                        if "t_tests" in self.fg_test_results:
                            fg_t_df = pd.DataFrame(
                                {
                                    "Variables": [
                                        f"{k[0]} vs {k[1]}" for k in self.fg_test_results["t_tests"]
                                    ],
                                    "t-statistique": [
                                        v[0] for v in self.fg_test_results["t_tests"].values()
                                    ],
                                    "p-value": [
                                        v[1] for v in self.fg_test_results["t_tests"].values()
                                    ],
                                    "Seuil alpha": [
                                        self.alpha for _ in self.fg_test_results["t_tests"].values()
                                    ],
                                    "Conclusion": [
                                        "Corrélation" if v[1] < self.alpha else "Pas de corrélation"
                                        for v in self.fg_test_results["t_tests"].values()
                                    ],
                                }
                            )
                            fg_t_df.to_excel(writer, sheet_name="Test_FG_t", index=False)

                    # Mix optimal
                    if self.include_stat_tests.get():
                        try:
                            x_vars = [x.strip() for x in self.x_vars.get().split("+")]
                            y_mean = np.mean(self.y)
                            x_means = [np.mean(self.data[var]) for var in x_vars]
                            coefficients = [self.model.params[var] for var in x_vars]

                            elasticities = [
                                coef * (x_mean / y_mean)
                                for coef, x_mean in zip(coefficients, x_means)
                            ]
                            sum_elasticities = sum(abs(e) for e in elasticities)
                            optimal_shares = [abs(e) / sum_elasticities for e in elasticities]

                            mix_df = pd.DataFrame(
                                {
                                    "Variable": x_vars,
                                    "Elasticité": elasticities,
                                    "Part_optimale": optimal_shares,
                                }
                            )
                            mix_df.to_excel(writer, sheet_name="Mix_optimal", index=False)
                        except:
                            pass

                messagebox.showinfo("Succès", f"Exporté vers {file_path}!")
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur d'exportation: {str(e)}")

    def export_to_word(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Aucun modèle à exporter!")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word", "*.docx"), ("Tous fichiers", "*.*")]
        )

        if file_path:
            try:
                doc = docx.Document()

                # Titre
                doc.add_heading("Résultats Économétriques", 0)
                doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
                doc.add_paragraph("\n")

                # Données
                if self.include_data.get():
                    doc.add_heading("Données", level=1)
                    doc.add_paragraph(f"Nombre d'observations: {len(self.data)}")
                    doc.add_paragraph(f"Variables: {', '.join(self.data.columns)}")
                    doc.add_paragraph("\n")

                # Résumé modèle
                if self.include_summary.get():
                    doc.add_heading("Résumé du Modèle", level=1)

                    stats = [
                        ("Variable dépendante", self.model.model.endog_names),
                        ("R²", f"{self.model.rsquared:.{self.decimal_places}f}"),
                        ("R² ajusté", f"{self.model.rsquared_adj:.{self.decimal_places}f}"),
                        ("F-statistique", f"{self.model.fvalue:.{self.decimal_places}f}"),
                        ("Prob(F-statistique)", f"{self.model.f_pvalue:.{self.decimal_places}f}"),
                        ("Log-vraisemblance", f"{self.model.llf:.{self.decimal_places}f}"),
                        ("AIC", f"{self.model.aic:.{self.decimal_places}f}"),
                        ("BIC", f"{self.model.bic:.{self.decimal_places}f}"),
                        ("Nb. observations", str(self.model.nobs)),
                    ]

                    table = doc.add_table(rows=1, cols=2)
                    table.style = "LightShading"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Statistique"
                    hdr_cells[1].text = "Valeur"

                    for stat, val in stats:
                        row_cells = table.add_row().cells
                        row_cells[0].text = stat
                        row_cells[1].text = val

                    doc.add_paragraph("\n")

                    # Coefficients
                    doc.add_heading("Coefficients", level=2)
                    coef_table = doc.add_table(rows=1, cols=5)
                    coef_table.style = "LightShading"
                    hdr_cells = coef_table.rows[0].cells
                    headers = ["Variable", "Coefficient", "Ecart-type", "t", "P>|t|"]

                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header

                    for i, var in enumerate(self.model.params.index):
                        row_cells = coef_table.add_row().cells
                        values = [
                            var,
                            f"{self.model.params[i]:.{self.decimal_places}f}",
                            f"{self.model.bse[i]:.{self.decimal_places}f}",
                            f"{self.model.tvalues[i]:.{self.decimal_places}f}",
                            f"{self.model.pvalues[i]:.{self.decimal_places}f}",
                        ]
                        for j, val in enumerate(values):
                            row_cells[j].text = val

                    doc.add_paragraph("\n")

                # Diagnostics
                if self.include_diagnostics.get():
                    doc.add_heading("Diagnostics", level=1)

                    # Tests
                    doc.add_heading("Tests de Diagnostic", level=2)
                    tests = [
                        (
                            "Durbin-Watson (autocorrélation)",
                            f"{durbin_watson(self.model.resid):.{self.decimal_places}f}",
                        ),
                        (
                            "Breusch-Pagan (hétéroscédasticité)",
                            f"LM = {het_breuschpagan(self.model.resid, self.model.model.exog)[0]:.{self.decimal_places}f}, "
                            + f"p-value = {het_breuschpagan(self.model.resid, self.model.model.exog)[1]:.{self.decimal_places}f}",
                        ),
                        (
                            "Anderson-Darling (normalité)",
                            f"A² = {normal_ad(self.model.resid)[0]:.{self.decimal_places}f}, "
                            + f"p-value = {normal_ad(self.model.resid)[1]:.{self.decimal_places}f}",
                        ),
                        (
                            "Jarque-Bera (normalité)",
                            f"JB = {jarque_bera(self.model.resid)[0]:.{self.decimal_places}f}, "
                            + f"p-value = {jarque_bera(self.model.resid)[1]:.{self.decimal_places}f}",
                        ),
                    ]

                    for test, result in tests:
                        doc.add_paragraph(f"{test}: {result}")

                    doc.add_paragraph("\n")

                # Multicolinéarité
                doc.add_heading("Multicolinéarité", level=1)

                # VIF
                doc.add_heading("Facteurs d'Inflation de la Variance (VIF)", level=2)
                vif_table = doc.add_table(rows=1, cols=2)
                vif_table.style = "LightShading"
                hdr_cells = vif_table.rows[0].cells
                hdr_cells[0].text = "Variable"
                hdr_cells[1].text = "VIF"

                vif_data = []
                for i, col in enumerate(self.X.columns):
                    if col == "const":
                        continue
                    vif = variance_inflation_factor(self.X.values, i)
                    vif_data.append((col, vif))

                vif_data.sort(key=lambda x: x[1], reverse=True)

                for var, vif in vif_data:
                    row_cells = vif_table.add_row().cells
                    row_cells[0].text = var
                    row_cells[1].text = f"{vif:.{self.decimal_places}f}"

                doc.add_paragraph("\n")

                # Test de Klein
                if self.include_stat_tests.get() and self.klein_test_results:
                    doc.add_heading("Test de Klein", level=2)
                    klein_table = doc.add_table(rows=1, cols=4)
                    klein_table.style = "LightShading"
                    hdr_cells = klein_table.rows[0].cells
                    hdr_cells[0].text = "Variables"
                    hdr_cells[1].text = "R² modèle"
                    hdr_cells[2].text = "r² variables"
                    hdr_cells[3].text = "Problème"

                    for (var1, var2), (r2, r2_vars) in self.klein_test_results.items():
                        row_cells = klein_table.add_row().cells
                        row_cells[0].text = f"{var1} vs {var2}"
                        row_cells[1].text = f"{r2:.4f}"
                        row_cells[2].text = f"{r2_vars:.4f}"
                        row_cells[3].text = "Oui" if r2 < r2_vars else "Non"

                    doc.add_paragraph("\n")

                # Test de Farrar-Glauber
                if self.include_stat_tests.get() and self.fg_test_results:
                    doc.add_heading("Test de Farrar-Glauber", level=2)

                    # Test global
                    doc.add_heading("Test global du Chi²", level=3)
                    fg_global_table = doc.add_table(rows=1, cols=4)
                    fg_global_table.style = "LightShading"
                    hdr_cells = fg_global_table.rows[0].cells
                    hdr_cells[0].text = "Test"
                    hdr_cells[1].text = "Statistique"
                    hdr_cells[2].text = "p-value"
                    hdr_cells[3].text = "Conclusion"

                    row_cells = fg_global_table.add_row().cells
                    row_cells[0].text = "Chi² global"
                    row_cells[1].text = f"{self.fg_test_results['global'][0]:.4f}"
                    row_cells[2].text = f"{self.fg_test_results['global'][2]:.4f}"
                    row_cells[3].text = (
                        "Multicolinéarité"
                        if self.fg_test_results["global"][2] < 0.05
                        else "Pas de multicolinéarité"
                    )

                    doc.add_paragraph("\n")

                    # Tests F
                    doc.add_heading("Tests F sur régressions auxiliaires", level=3)
                    fg_f_table = doc.add_table(rows=1, cols=4)
                    fg_f_table.style = "LightShading"
                    hdr_cells = fg_f_table.rows[0].cells
                    hdr_cells[0].text = "Variable"
                    hdr_cells[1].text = "F-statistique"
                    hdr_cells[2].text = "p-value"
                    hdr_cells[3].text = "Conclusion"

                    for var, (f_val, f_pval) in self.fg_test_results["f_tests"].items():
                        row_cells = fg_f_table.add_row().cells
                        row_cells[0].text = var
                        row_cells[1].text = f"{f_val:.4f}"
                        row_cells[2].text = f"{f_pval:.4f}"
                        row_cells[3].text = "Liée aux autres" if f_pval < 0.05 else "Non liée"

                    doc.add_paragraph("\n")

                    # Tests t
                    if "t_tests" in self.fg_test_results:
                        doc.add_heading("Tests t sur corrélations partielles", level=3)
                        fg_t_table = doc.add_table(rows=1, cols=4)
                        fg_t_table.style = "LightShading"
                        hdr_cells = fg_t_table.rows[0].cells
                        hdr_cells[0].text = "Variables"
                        hdr_cells[1].text = "t-statistique"
                        hdr_cells[2].text = "p-value"
                        hdr_cells[3].text = "Conclusion"

                        for (var1, var2), (t_val, t_pval) in self.fg_test_results[
                            "t_tests"
                        ].items():
                            row_cells = fg_t_table.add_row().cells
                            row_cells[0].text = f"{var1} vs {var2}"
                            row_cells[1].text = f"{t_val:.4f}"
                            row_cells[2].text = f"{t_pval:.4f}"
                            row_cells[3].text = (
                                "Corrélation" if t_pval < 0.05 else "Pas de corrélation"
                            )

                        doc.add_paragraph("\n")

                # Corrélations partielles
                if self.include_partial_corr.get() and self.partial_corr_results:
                    doc.add_heading("Corrélations Partielles", level=1)
                    partial_table = doc.add_table(rows=1, cols=3)
                    partial_table.style = "LightShading"
                    hdr_cells = partial_table.rows[0].cells
                    hdr_cells[0].text = "Variable 1"
                    hdr_cells[1].text = "Variable 2"
                    hdr_cells[2].text = "Corrélation Partielle"

                    for (var1, var2), corr in self.partial_corr_results.items():
                        row_cells = partial_table.add_row().cells
                        row_cells[0].text = var1
                        row_cells[1].text = var2
                        row_cells[2].text = f"{corr:.{self.decimal_places}f}"

                    doc.add_paragraph("\n")

                # Mix optimal
                if self.include_stat_tests.get():
                    try:
                        x_vars = [x.strip() for x in self.x_vars.get().split("+")]
                        y_mean = np.mean(self.y)
                        x_means = [np.mean(self.data[var]) for var in x_vars]
                        coefficients = [self.model.params[var] for var in x_vars]

                        elasticities = [
                            coef * (x_mean / y_mean) for coef, x_mean in zip(coefficients, x_means)
                        ]
                        sum_elasticities = sum(abs(e) for e in elasticities)
                        optimal_shares = [abs(e) / sum_elasticities for e in elasticities]

                        doc.add_heading("Mix Optimal", level=1)
                        mix_table = doc.add_table(rows=1, cols=3)
                        mix_table.style = "LightShading"
                        hdr_cells = mix_table.rows[0].cells
                        hdr_cells[0].text = "Variable"
                        hdr_cells[1].text = "Elasticité"
                        hdr_cells[2].text = "Part optimale"

                        for var, elasticity, share in zip(x_vars, elasticities, optimal_shares):
                            row_cells = mix_table.add_row().cells
                            row_cells[0].text = var
                            row_cells[1].text = f"{elasticity:.4f}"
                            row_cells[2].text = f"{share * 100:.2f}%"

                        doc.add_paragraph("\n")
                    except:
                        pass

                # Interprétations
                if self.include_interpretations.get():
                    doc.add_heading("Interprétations", level=1)
                    interpretations = self.get_interpretations().split("\n")
                    for interp in interpretations:
                        if interp.strip():
                            doc.add_paragraph(interp)
                    doc.add_paragraph("\n")

                # Formules
                if self.include_formulas.get():
                    doc.add_heading("Formules Mathématiques", level=1)
                    formulas = self.get_model_formulas().split("\n")
                    for formula in formulas:
                        if formula.strip():
                            p = doc.add_paragraph()
                            p.add_run(formula).font.name = "Courier New"
                    doc.add_paragraph("\n")

                # Étapes de calcul
                if self.include_calculation_steps.get():
                    doc.add_heading("Étapes de Calcul", level=1)
                    steps = self.get_calculation_steps().split("\n")
                    for step in steps:
                        if step.strip():
                            p = doc.add_paragraph()
                            p.add_run(step).font.name = "Courier New"
                    doc.add_paragraph("\n")

                # Hypothèses des tests
                if self.include_hypothesis.get():
                    doc.add_heading("Hypothèses des Tests", level=1)

                    # Test t
                    doc.add_heading("Test t de significativité", level=2)
                    doc.add_paragraph(
                        "H0: Le coefficient est égal à 0\n"
                        "H1: Le coefficient est différent de 0\n"
                        "Règle de décision: Rejet de H0 si p-value < 0.05"
                    )
                    doc.add_paragraph("\n")

                    # Test F
                    doc.add_heading("Test F global", level=2)
                    doc.add_paragraph(
                        "H0: Tous les coefficients (sauf constante) = 0\n"
                        "H1: Au moins un coefficient ≠ 0\n"
                        "Règle de décision: Rejet de H0 si p-value < 0.05"
                    )
                    doc.add_paragraph("\n")

                    # Test de Klein
                    doc.add_heading("Test de Klein", level=2)
                    doc.add_paragraph(
                        "H0: Pas de multicolinéarité problématique\n"
                        "H1: Présence de multicolinéarité problématique\n"
                        "Règle: Si R² modèle < r² entre deux variables ⇒ multicolinéarité"
                    )
                    doc.add_paragraph("\n")

                    # Test de Farrar-Glauber
                    doc.add_heading("Test de Farrar-Glauber", level=2)
                    doc.add_paragraph(
                        "1. Test global du Chi²:\n"
                        "   H0: Matrice de corrélation = Matrice identité\n"
                        "   H1: Matrice de corrélation ≠ Matrice identité\n"
                        "   Règle: Rejet de H0 si p-value < 0.05\n\n"
                        "2. Tests F sur régressions auxiliaires:\n"
                        "   H0: Tous les coefficients = 0\n"
                        "   H1: Au moins un coefficient ≠ 0\n"
                        "   Règle: Rejet de H0 si p-value < 0.05\n\n"
                        "3. Tests t sur corrélations partielles:\n"
                        "   H0: ρ(Xi,Xj|autres) = 0\n"
                        "   H1: ρ(Xi,Xj|autres) ≠ 0\n"
                        "   Règle: Rejet de H0 si p-value < 0.05"
                    )
                    doc.add_paragraph("\n")

                doc.save(file_path)
                messagebox.showinfo("Succès", f"Exporté vers {file_path}!")
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur d'exportation: {str(e)}")

    def export_to_pdf(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Aucun modèle à exporter!")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf"), ("Tous fichiers", "*.*")],
            title="Enregistrer le rapport PDF",
        )

        if not file_path:  # L'utilisateur a annulé
            return

        try:
            from datetime import datetime

            from fpdf import FPDF
            from statsmodels.stats.diagnostic import het_breuschpagan, normal_ad
            from statsmodels.stats.outliers_influence import variance_inflation_factor
            from statsmodels.stats.stattools import durbin_watson, jarque_bera

            # Création du PDF en format paysage
            pdf = FPDF(orientation="L")  # L pour landscape (paysage)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=10)

            # Largeur utile en paysage (297mm - marges)
            usable_width = 280

            # Titre
            pdf.set_font("Arial", "B", 14)
            pdf.cell(usable_width, 10, txt="Résultats Économétriques", ln=1, align="C")
            pdf.set_font("Arial", "", 10)
            pdf.cell(
                usable_width,
                8,
                txt=f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
                ln=1,
                align="C",
            )
            pdf.ln(5)

            try:
                # Données
                if self.include_data.get() and hasattr(self, "data"):
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(usable_width, 8, txt="Données", ln=1)
                    pdf.set_font("Arial", "", 10)

                    # Tableau des données
                    col_widths = [60, 220]  # Colonnes pour les descriptions et valeurs

                    pdf.cell(col_widths[0], 8, txt="Nombre d'observations:", border=1)
                    pdf.cell(col_widths[1], 8, txt=str(len(self.data)), border=1, ln=1)

                    if hasattr(self.data, "columns"):
                        pdf.cell(col_widths[0], 8, txt="Variables:", border=1)
                        vars_text = ", ".join(self.data.columns)
                        # Gestion des longues listes de variables
                        if len(vars_text) > 100:
                            vars_text = vars_text[:100] + "..."
                        pdf.cell(col_widths[1], 8, txt=vars_text, border=1, ln=1)

                    pdf.ln(5)

                # Résumé modèle
                if self.include_summary.get() and hasattr(self.model, "model"):
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(usable_width, 8, txt="Résumé du Modèle", ln=1)
                    pdf.set_font("Arial", "", 10)

                    stats = [
                        ("Variable dépendante", getattr(self.model.model, "endog_names", "N/A")),
                        ("R²", f"{getattr(self.model, 'rsquared', 0):.{self.decimal_places}f}"),
                        (
                            "R² ajusté",
                            f"{getattr(self.model, 'rsquared_adj', 0):.{self.decimal_places}f}",
                        ),
                        (
                            "F-statistique",
                            f"{getattr(self.model, 'fvalue', 0):.{self.decimal_places}f}",
                        ),
                        (
                            "Prob(F-statistique)",
                            f"{getattr(self.model, 'f_pvalue', 0):.{self.decimal_places}f}",
                        ),
                        (
                            "Log-vraisemblance",
                            f"{getattr(self.model, 'llf', 0):.{self.decimal_places}f}",
                        ),
                        ("AIC", f"{getattr(self.model, 'aic', 0):.{self.decimal_places}f}"),
                        ("BIC", f"{getattr(self.model, 'bic', 0):.{self.decimal_places}f}"),
                        ("Nb. observations", str(getattr(self.model, "nobs", 0))),
                    ]

                    # Tableau sur toute la largeur
                    col_width = usable_width / 2
                    line_height = 8

                    for stat, val in stats:
                        pdf.cell(col_width, line_height, txt=stat, border=1)
                        pdf.cell(col_width, line_height, txt=str(val), border=1, ln=1)

                    pdf.ln(5)

                    # Coefficients - version optimisée pour le paysage
                    if hasattr(self.model, "params"):
                        pdf.set_font("Arial", "B", 12)
                        pdf.cell(usable_width, 8, txt="Coefficients", ln=1)
                        pdf.set_font("Arial", "", 10)

                        # Largeurs de colonnes proportionnelles à la largeur utilisable
                        headers = ["Variable", "Coefficient", "Ecart-type", "t", "P>|t|"]
                        col_widths = [
                            usable_width * 0.25,  # Variable
                            usable_width * 0.20,  # Coef
                            usable_width * 0.20,  # Std err
                            usable_width * 0.15,  # t
                            usable_width * 0.20,  # p-value
                        ]

                        # En-têtes
                        for i, header in enumerate(headers):
                            pdf.cell(col_widths[i], line_height, txt=header, border=1, align="C")
                        pdf.ln()

                        # Données
                        for i, var in enumerate(self.model.params.index):
                            values = [
                                str(var),
                                f"{self.model.params[i]:.{self.decimal_places}f}",
                                f"{getattr(self.model, 'bse', [0] * len(self.model.params))[i]:.{self.decimal_places}f}",
                                f"{getattr(self.model, 'tvalues', [0] * len(self.model.params))[i]:.{self.decimal_places}f}",
                                f"{getattr(self.model, 'pvalues', [0] * len(self.model.params))[i]:.{self.decimal_places}f}",
                            ]
                            for j, val in enumerate(values):
                                align = "R" if j > 0 else "L"  # Alignement droite pour les nombres
                                pdf.cell(col_widths[j], line_height, txt=val, border=1, align=align)
                            pdf.ln()

                        pdf.ln(5)

                # Diagnostics - version compacte
                if self.include_diagnostics.get() and hasattr(self.model, "resid"):
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(usable_width, 8, txt="Diagnostics", ln=1)
                    pdf.set_font("Arial", "", 10)

                    # Obtenir les résultats avec gestion des erreurs
                    try:
                        dw = durbin_watson(self.model.resid)
                    except:
                        dw = "N/A"

                    try:
                        bp_test = het_breuschpagan(self.model.resid, self.model.model.exog)
                    except:
                        bp_test = ("N/A", "N/A")

                    try:
                        ad_test = normal_ad(self.model.resid)
                    except:
                        ad_test = ("N/A", "N/A")

                    try:
                        jb_test = jarque_bera(self.model.resid)
                    except:
                        jb_test = ("N/A", "N/A")

                    tests = [
                        (
                            "Durbin-Watson (autocorrélation)",
                            f"{dw:.{self.decimal_places}f}"
                            if isinstance(dw, (int, float))
                            else str(dw),
                        ),
                        (
                            "Breusch-Pagan (hétéroscédasticité)",
                            f"LM={bp_test[0]:.{self.decimal_places}f}, p={bp_test[1]:.{self.decimal_places}f}"
                            if all(isinstance(x, (int, float)) for x in bp_test)
                            else "N/A",
                        ),
                        (
                            "Anderson-Darling (normalité)",
                            f"A²={ad_test[0]:.{self.decimal_places}f}, p={ad_test[1]:.{self.decimal_places}f}"
                            if all(isinstance(x, (int, float)) for x in ad_test)
                            else "N/A",
                        ),
                        (
                            "Jarque-Bera (normalité)",
                            f"JB={jb_test[0]:.{self.decimal_places}f}, p={jb_test[1]:.{self.decimal_places}f}"
                            if all(isinstance(x, (int, float)) for x in jb_test)
                            else "N/A",
                        ),
                    ]

                    # Tableau sur 2 colonnes
                    col_width = usable_width / 2
                    for i in range(0, len(tests), 2):
                        # Première colonne
                        test1, result1 = tests[i]
                        pdf.cell(col_width, 8, txt=test1, border=1)
                        pdf.cell(col_width, 8, txt=result1, border=1, ln=1)

                        # Deuxième colonne (si existe)
                        if i + 1 < len(tests):
                            test2, result2 = tests[i + 1]
                            pdf.cell(col_width, 8, txt=test2, border=1)
                            pdf.cell(col_width, 8, txt=result2, border=1, ln=1)

                    pdf.ln(5)

                # Multicolinéarité - version optimisée
                if hasattr(self, "X") and hasattr(self.X, "columns"):
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(usable_width, 8, txt="Multicolinéarité", ln=1)
                    pdf.set_font("Arial", "", 10)

                    # VIF
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(usable_width, 8, txt="Facteurs d'Inflation de la Variance (VIF)", ln=1)
                    pdf.set_font("Arial", "", 10)

                    vif_data = []
                    for i, col in enumerate(self.X.columns):
                        if col == "const":
                            continue
                        try:
                            vif = variance_inflation_factor(self.X.values, i)
                            vif_data.append((col, vif))
                        except:
                            vif_data.append((col, "N/A"))

                    if vif_data:
                        vif_data.sort(
                            key=lambda x: x[1] if isinstance(x[1], (int, float)) else 0,
                            reverse=True,
                        )

                        # Tableau sur 2 colonnes pour économiser de l'espace
                        col_width = usable_width / 2
                        line_height = 8

                        for i in range(0, len(vif_data), 2):
                            # Première colonne
                            var1, vif1 = vif_data[i]
                            pdf.cell(col_width * 0.6, line_height, txt=var1, border=1)
                            pdf.cell(
                                col_width * 0.4,
                                line_height,
                                txt=f"{vif1:.{self.decimal_places}f}"
                                if isinstance(vif1, (int, float))
                                else str(vif1),
                                border=1,
                            )

                            # Deuxième colonne (si existe)
                            if i + 1 < len(vif_data):
                                var2, vif2 = vif_data[i + 1]
                                pdf.cell(col_width * 0.6, line_height, txt=var2, border=1)
                                pdf.cell(
                                    col_width * 0.4,
                                    line_height,
                                    txt=f"{vif2:.{self.decimal_places}f}"
                                    if isinstance(vif2, (int, float))
                                    else str(vif2),
                                    border=1,
                                )

                            pdf.ln()

                        pdf.ln(5)

                # Autres sections (ANOVA, etc.) avec le même format...

            except Exception as e:
                pdf.set_font("Arial", size=10)
                pdf.multi_cell(
                    usable_width, 8, txt=f"Erreur lors de la génération du rapport: {str(e)}", ln=1
                )

            try:
                pdf.output(file_path)
                messagebox.showinfo("Succès", f"Rapport exporté avec succès vers:\n{file_path}")
            except Exception as e:
                messagebox.showerror(
                    "Erreur", f"Impossible d'enregistrer le fichier PDF:\n{str(e)}"
                )

        except ImportError as e:
            messagebox.showerror(
                "Erreur",
                f"Bibliothèque manquante: {str(e)}\nVeuillez installer les dépendances nécessaires.",
            )
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur inattendue lors de l'export PDF:\n{str(e)}")

    def export_plots(self):
        if not self.model:
            messagebox.showwarning("Avertissement", "Aucun modèle à exporter!")
            return

        folder_path = filedialog.askdirectory()
        if folder_path:
            try:
                # Résidus vs ajustés
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                ax.scatter(self.model.fittedvalues, self.model.resid)
                ax.axhline(y=0, color="r", linestyle="--")
                ax.set_xlabel("Valeurs ajustées")
                ax.set_ylabel("Résidus")
                ax.set_title("Résidus vs Valeurs Ajustées")
                plt.savefig(os.path.join(folder_path, "residus_vs_ajustes.png"))

                # QQ plot
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                sm.qqplot(self.model.resid, line="s", ax=ax)
                ax.set_title("QQ Plot des Résidus")
                plt.savefig(os.path.join(folder_path, "qqplot.png"))

                # Matrice de corrélation
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                corr = self.data.corr()
                sns.heatmap(corr, annot=True, fmt=f".{self.decimal_places}f", ax=ax)
                ax.set_title("Matrice de Corrélation")
                plt.savefig(os.path.join(folder_path, "correlation.png"))

                # Leverage plot
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                sm.graphics.influence_plot(self.model, ax=ax)
                ax.set_title("Leverage Plot")
                plt.savefig(os.path.join(folder_path, "leverage.png"))

                # Autocorrélation
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                plot_acf(self.model.resid, ax=ax)
                ax.set_title("Fonction d'Autocorrélation")
                plt.savefig(os.path.join(folder_path, "autocorrelation.png"))

                # Autocorrélation partielle
                self.figure.clf()
                ax = self.figure.add_subplot(111)
                plot_pacf(self.model.resid, ax=ax)
                ax.set_title("Fonction d'Autocorrélation Partielle")
                plt.savefig(os.path.join(folder_path, "autocorrelation_partielle.png"))

                messagebox.showinfo("Succès", f"Graphiques exportés dans {folder_path}!")
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur d'exportation: {str(e)}")


def main() -> None:
    """Point d'entree principal de l'application econometrique."""
    root = tk.Tk()
    EconometrieAvanceeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
