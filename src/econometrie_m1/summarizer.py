"""Application de resume automatique de documents."""

from __future__ import annotations

import os
import re
import textwrap
import threading
import tkinter as tk
from tkinter import Text, filedialog, messagebox, ttk

import customtkinter as ctk
import docx
import pdfplumber
from fpdf import FPDF
from transformers import AutoTokenizer, pipeline

# ──────────────────────────────────────────────────────────────
# 1.  MODELS DECLARATION
# ──────────────────────────────────────────────────────────────
SMALL = "sshleifer/distilbart-cnn-12-6"
LARGE = "pszemraj/led-large-book-summary"

summ_small = pipeline("summarization", model=SMALL)
summ_large = pipeline("summarization", model=LARGE)

tok_small = AutoTokenizer.from_pretrained(SMALL)
tok_large = AutoTokenizer.from_pretrained(LARGE)

trans_en_fr = pipeline("translation", model="Helsinki-NLP/opus-mt-en-fr")
trans_fr_en = pipeline("translation", model="Helsinki-NLP/opus-mt-fr-en")

# ──────────────────────────────────────────────────────────────
# 2.  HELPERS
# ──────────────────────────────────────────────────────────────


def clean(txt: str) -> str:
    txt = re.sub(r"https?://\\S+", "", txt)
    txt = re.sub(r"\\s+", " ", txt)
    return txt.strip()


def shorten(path: str, n: int = 80) -> str:
    return path if len(path) <= n else "…" + path[-n:]


def get_text(path: str) -> str:
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return clean("\n".join(p.extract_text() or "" for p in pdf.pages))
    if path.endswith(".docx"):
        doc = docx.Document(path)
        return clean("\n".join(p.text for p in doc.paragraphs))
    if path.endswith(".txt"):
        with open(path, encoding="utf-8") as f:
            return clean(f.read())
    messagebox.showerror("Erreur", "Format non pris en charge")
    return ""


def choose_pipe(txt: str):
    return (
        (summ_small, tok_small, 1024)
        if len(tok_small.encode(txt)) <= 1024
        else (summ_large, tok_large, 4096)
    )


def chunk(txt: str, tokenizer, max_tok: int):
    chunks, cur, c_len = [], "", 0
    for para in txt.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p_len = len(tokenizer.encode(para)) + 1
        if c_len + p_len > max_tok:
            if cur:
                chunks.append(cur.strip())
            cur, c_len = para, p_len
        else:
            cur += "\n\n" + para
            c_len += p_len
    if cur:
        chunks.append(cur.strip())
    return chunks


# ──────────────────────────────────────────────────────────────
# 3.  THEME & STATE
# ──────────────────────────────────────────────────────────────
COL_BG = "#F1F5F9"  # slate‑50
COL_TXT = "#0F172A"  # slate‑900
COL_PRIM = "#2563EB"  # blue‑600  (Résumer)
COL_PRIM_H = "#1D4ED8"  # blue‑700
COL_SEC = "#F97316"  # orange‑500 (Ouvrir)
COL_SEC_H = "#EA580C"  # orange‑600
COL_ACC = "#10B981"  # emerald‑500 (Traduire)
COL_ACC_H = "#059669"  # emerald‑600

FILE = ""
HIST = []
HIST_PATH = "summary_history.txt"

# ──────────────────────────────────────────────────────────────
# 4.  UI ROOT (CustomTkinter)
# ──────────────────────────────────────────────────────────────
# ──────────────────────────────────────────────────────────────
# 4.  UI ROOT (CustomTkinter)
# ──────────────────────────────────────────────────────────────
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

root = ctk.CTk()
root.title("Résumé Automatique de Documents")
root.geometry("1240x720")
root.configure(fg_color=COL_BG)

sty = ttk.Style(root)
sty.theme_use("clam")
sty.configure("TLabel", background=COL_BG, foreground=COL_TXT, font=("Segoe UI", 10))
sty.configure("Title.TLabel", font=("Segoe UI", 22, "bold"))


def round_button(master, text, command):
    return ctk.CTkButton(
        master,
        text=text,
        command=command,
        corner_radius=18,
        font=("Segoe UI", 10, "bold"),
        width=120,
        height=34,
    )


def colored_button(master, text, command, color, hover):
    return ctk.CTkButton(
        master,
        text=text,
        command=command,
        fg_color=color,
        hover_color=hover,
        text_color="white",
        corner_radius=18,
        font=("Segoe UI", 10, "bold"),
        width=120,
        height=34,
    )


# ─── Header ────────────────────────────────
header = ttk.Label(root, text="Résumé Automatique de Documents", style="Title.TLabel")
header.grid(row=0, column=0, columnspan=5, pady=(14, 6), sticky="w", padx=14)

path_lbl = ttk.Label(root, text="Aucun fichier sélectionné", anchor="w", wraplength=850)
path_lbl.grid(row=1, column=0, columnspan=5, sticky="w", padx=14)

# ─── Row: Fichier + Résumer + Slider ────────
open_btn = colored_button(root, "Ouvrir", lambda: on_open(), COL_SEC, COL_SEC_H)
run_btn = colored_button(root, "Résumer", lambda: run_summary(), COL_PRIM, COL_PRIM_H)

length_var = tk.IntVar(value=150)
len_lbl = ttk.Label(root, text="Taille max (tokens)")
slider = tk.Scale(
    root,
    from_=50,
    to=800,
    variable=length_var,
    orient="horizontal",
    length=200,
    resolution=10,
    tickinterval=150,
    bg=COL_BG,
    highlightthickness=0,
)
len_val = ttk.Label(root, textvariable=length_var)

open_btn.grid(row=2, column=0, padx=14, pady=6, sticky="ew")
run_btn.grid(row=2, column=1, padx=6, pady=6, sticky="ew")
len_lbl.grid(row=2, column=2, sticky="w")
slider.grid(row=2, column=3, padx=6, sticky="w")
len_val.grid(row=2, column=4, sticky="w")

# ─── Progress Bars ─────────────────────────
prog = ttk.Progressbar(root, mode="determinate")
prog.grid(row=3, column=0, columnspan=5, sticky="ew", padx=14)
prog.grid_remove()

prog_tr = ttk.Progressbar(root, mode="indeterminate")
prog_tr.grid(row=3, column=0, columnspan=5, sticky="ew", padx=14)
prog_tr.grid_remove()

# ─── Résumé Result Area ────────────────────
res_txt = Text(root, wrap=tk.WORD, height=15, font=("Segoe UI", 10))
res_txt.grid(row=3, column=0, columnspan=5, padx=14, pady=(10, 6), sticky="nsew")

# ─── Traduction + Export (sur la même ligne) ───
bottom_frame = ttk.Frame(root, style="TFrame")
bottom_frame.grid(row=4, column=0, columnspan=5, sticky="ew", padx=14, pady=(0, 4))

lang = tk.StringVar(value="Français")
lang_sel = ttk.OptionMenu(bottom_frame, lang, "Français", "Français", "Anglais")
tr_btn = round_button(bottom_frame, "Traduire", lambda: run_translate())
pdf_btn = round_button(bottom_frame, "PDF", lambda: save_pdf())
wx_btn = round_button(bottom_frame, "DOCX", lambda: save_docx())

lang_sel.pack(side=tk.LEFT, padx=(0, 8))
tr_btn.pack(side=tk.LEFT, padx=(0, 20))
pdf_btn.pack(side=tk.RIGHT, padx=4)
wx_btn.pack(side=tk.RIGHT, padx=4)

# ─── Historique à droite ────────────────────
hist_lbl = ttk.Label(root, text="Historique", style="Title.TLabel")
hist_lbl.grid(row=0, column=5, padx=8, sticky="w")

scroll = ttk.Scrollbar(root, orient="vertical")
hist_txt = Text(root, width=55, wrap=tk.WORD, yscrollcommand=scroll.set, font=("Segoe UI", 10))
scroll.config(command=hist_txt.yview)
hist_txt.grid(row=1, column=5, rowspan=6, pady=8, padx=(0, 8), sticky="nsew")
scroll.grid(row=1, column=6, rowspan=6, sticky="ns")

# ─── Layout config ──────────────────────────
for c in range(6):
    root.grid_columnconfigure(c, weight=1)
root.grid_rowconfigure(4, weight=1)  # résumé


# ——————————————————————————————————————————————
# 5.  LOGIC
# ——————————————————————————————————————————————


def history_refresh():
    hist_txt.config(state=tk.NORMAL)
    hist_txt.delete(1.0, tk.END)
    hist_txt.tag_configure("wrap", lmargin1=6, lmargin2=6, spacing3=4)
    sep = "─" * 110 + "\n"
    for title, summ, trans in HIST:
        hist_txt.insert(tk.END, title + "\n", "wrap")
        hist_txt.insert(tk.END, "Résumé : " + textwrap.fill(summ, 100) + "\n", "wrap")
        if trans:
            hist_txt.insert(tk.END, "Traduction : " + textwrap.fill(trans, 100) + "\n", "wrap")
        hist_txt.insert(tk.END, sep, "wrap")
    hist_txt.config(state=tk.DISABLED)


def hist_load():
    if not os.path.exists(HIST_PATH):
        return
    with open(HIST_PATH, encoding="utf-8") as f:
        for line in f:
            parts = line.strip().split("::")
            if len(parts) == 3:
                HIST.append(tuple(parts))
    history_refresh()


def hist_save(title, summ, trans=""):
    with open(HIST_PATH, "a", encoding="utf-8") as f:
        f.write(f"{title}::{summ}::{trans}\n")


# ─── Callbacks ──────────────────────────────


def on_open():
    global FILE
    FILE = filedialog.askopenfilename(
        title="Choisir un fichier", filetypes=(("Documents", "*.pdf *.docx *.txt"),)
    )
    if FILE:
        path_lbl.configure(text=f"Fichier : {shorten(FILE)}")


def run_summary():
    if not FILE:
        messagebox.showwarning("Choix manquant", "Veuillez sélectionner un fichier")
        return

    title = os.path.splitext(os.path.basename(FILE))[0]
    txt = get_text(FILE)
    if not txt:
        return

    max_len = length_var.get()
    min_len = max(30, int(max_len * 0.4))

    # Nettoyage optionnel
    txt = clean(txt)

    # Teste avec les deux tokenizers pour trouver la meilleure stratégie
    if len(tok_large.encode(txt)) > 16384:
        messagebox.showwarning(
            "Texte très long",
            "Ce document est extrêmement long. Il sera traité en plusieurs blocs.",
        )

    pipe, tok, max_tok = choose_pipe(txt)
    chunks = chunk(txt, tok, max_tok)

    if not chunks:
        messagebox.showerror("Erreur", "Impossible de découper le texte pour résumé.")
        return

    prog.grid()
    prog["maximum"], prog["value"] = len(chunks), 0
    res_txt.delete(1.0, tk.END)
    parts = []

    try:
        for i, ch in enumerate(chunks, 1):
            result = pipe(
                ch, max_length=max_len, min_length=min_len, do_sample=False, truncation=True
            )
            parts.append(result[0]["summary_text"])
            prog["value"] = i
            root.update_idletasks()

        summ = " ".join(parts)
        res_txt.insert(tk.END, summ)
        HIST.append((title, summ, ""))
        hist_save(title, summ)
        history_refresh()

    except Exception as e:
        messagebox.showerror("Erreur", str(e))
    finally:
        prog.grid_remove()


def run_translate():
    txt = res_txt.get(1.0, tk.END).strip()
    if not txt:
        return

    # ── Choix du sens de traduction ──────────
    model = trans_en_fr if lang.get() == "Français" else trans_fr_en

    # ── Animation indéterminée ───────────────
    def _translate():
        try:
            trans = model(txt, max_length=512)[0]["translation_text"]
            res_txt.delete(1.0, tk.END)
            res_txt.insert(tk.END, trans)

            if HIST:
                title, summ, _ = HIST[-1]
                HIST[-1] = (title, summ, trans)
                hist_save(title, summ, trans)
                history_refresh()
        finally:
            prog_tr.stop()
            prog_tr.grid_remove()

    prog_tr.grid()
    prog_tr.start()
    threading.Thread(target=_translate, daemon=True).start()


def save_docx():
    txt = res_txt.get(1.0, tk.END).strip()
    if not txt:
        return
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=(("DOCX", "*.docx"),))
    if path:
        d = docx.Document()
        d.add_paragraph(txt)
        d.save(path)
        messagebox.showinfo("OK", "DOCX enregistré")


def save_pdf():
    txt = res_txt.get(1.0, tk.END).strip()
    if not txt:
        return
    path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=(("PDF", "*.pdf"),))
    if path:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt)
        pdf.output(path)
        messagebox.showinfo("OK", "PDF enregistré")


# ─── Bindings & Init ────────────────────────
open_btn.configure(command=on_open)
run_btn.configure(command=run_summary)
tr_btn.configure(command=run_translate)
wx_btn.configure(command=save_docx)
pdf_btn.configure(command=save_pdf)

# Poids des colonnes / lignes pour le redimensionnement
for c in range(7):
    root.grid_columnconfigure(c, weight=1)
root.grid_rowconfigure(4, weight=1)

hist_load()
root.mainloop()


def main() -> None:
    """Point d'entree principal de l'application de resume."""
    hist_load()
    root.mainloop()
