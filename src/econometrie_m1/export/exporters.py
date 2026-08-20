from __future__ import annotations

from typing import Any

import numpy as np
import pandas as pd
from statsmodels.regression.linear_model import RegressionResultsWrapper
from statsmodels.stats.diagnostic import het_breuschpagan, normal_ad
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.stattools import durbin_watson, jarque_bera


def export_results_to_excel(
    file_path: str,
    model: RegressionResultsWrapper,
    data: pd.DataFrame,
    X: pd.DataFrame,
    y: pd.Series,
    alpha: float = 0.05,
    decimal_places: int = 4,
    include_data: bool = True,
    include_summary: bool = True,
    include_matrices: bool = False,
    include_diagnostics: bool = False,
    include_partial_corr: bool = False,
    include_stat_tests: bool = False,
    partial_corr_results: dict[tuple[str, str], float] | None = None,
    klein_test_results: dict[tuple[str, str], tuple[float, float]] | None = None,
    fg_test_results: dict[str, Any] | None = None,
    x_var_names: list[str] | None = None,
) -> None:
    """Exporte les resultats de l'analyse econometrique vers Excel."""
    with pd.ExcelWriter(file_path) as writer:
        if include_data:
            data.to_excel(writer, sheet_name="Données")

        if include_summary:
            summary_df = pd.DataFrame(
                {
                    "Statistique": [
                        "R²",
                        "R² ajusté",
                        "F-statistique",
                        "Prob(F-statistique)",
                        "Log-vraisemblance",
                        "AIC",
                        "BIC",
                        "Nb. Observations",
                        "Seuil alpha",
                    ],
                    "Valeur": [
                        model.rsquared,
                        model.rsquared_adj,
                        model.fvalue,
                        model.f_pvalue,
                        model.llf,
                        model.aic,
                        model.bic,
                        model.nobs,
                        alpha,
                    ],
                }
            )
            summary_df.to_excel(writer, sheet_name="Résumé", index=False)

            params_df = pd.DataFrame(
                {
                    "Variable": model.params.index,
                    "Coefficient": model.params.values,
                    "Ecart-type": model.bse.values,
                    "t": model.tvalues.values,
                    "P>|t|": model.pvalues.values,
                    "Significatif": ["Oui" if p < alpha else "Non" for p in model.pvalues.values],
                }
            )
            params_df.to_excel(writer, sheet_name="Coefficients", index=False)

        if include_matrices:
            X_df = pd.DataFrame(X)
            X_df.to_excel(writer, sheet_name="Matrice_X")

            xtx = np.dot(X.T, X)
            pd.DataFrame(xtx, index=X.columns, columns=X.columns).to_excel(
                writer, sheet_name="Matrice_XtX"
            )

            xty = np.dot(X.T, y)
            pd.DataFrame(xty, index=X.columns, columns=["X'Y"]).to_excel(
                writer, sheet_name="Matrice_XtY"
            )

            try:
                xtx_inv = np.linalg.inv(xtx)
                pd.DataFrame(xtx_inv, index=X.columns, columns=X.columns).to_excel(
                    writer, sheet_name="Matrice_XtX_inv"
                )
            except np.linalg.LinAlgError:
                pass

            pd.DataFrame(model.cov_params()).to_excel(writer, sheet_name="Matrice_VarCov")

        if include_diagnostics:
            resid = model.resid
            exog = model.model.exog

            tests_df = pd.DataFrame(
                {
                    "Test": ["Durbin-Watson", "Breusch-Pagan", "Anderson-Darling", "Jarque-Bera"],
                    "Statistique": [
                        durbin_watson(resid),
                        het_breuschpagan(resid, exog)[0],
                        normal_ad(resid)[0],
                        jarque_bera(resid)[0],
                    ],
                    "p-value": [
                        "",
                        het_breuschpagan(resid, exog)[1],
                        normal_ad(resid)[1],
                        jarque_bera(resid)[1],
                    ],
                    "Seuil alpha": [alpha] * 4,
                    "Conclusion": [
                        "",
                        "Hétéroscédasticité" if het_breuschpagan(resid, exog)[1] < alpha else "OK",
                        "Non-normalité" if normal_ad(resid)[1] < alpha else "OK",
                        "Non-normalité" if jarque_bera(resid)[1] < alpha else "OK",
                    ],
                }
            )
            tests_df.to_excel(writer, sheet_name="Tests_Diagnostic", index=False)

            resid_df = pd.DataFrame(
                {
                    "Y_observé": y,
                    "Y_prédit": model.fittedvalues,
                    "Résidu": resid,
                    "Résidu_standardisé": model.get_influence().resid_studentized_internal,
                    "Leverage": model.get_influence().hat_matrix_diag,
                }
            )
            resid_df.to_excel(writer, sheet_name="Résidus", index=False)

        vif_data = []
        for i, col in enumerate(X.columns):
            if col == "const":
                continue
            vif = variance_inflation_factor(X.values, i)
            vif_data.append((col, vif))

        pd.DataFrame(vif_data, columns=["Variable", "VIF"]).to_excel(
            writer, sheet_name="Multicolinearite", index=False
        )

        if include_partial_corr and partial_corr_results:
            partial_corr_df = pd.DataFrame(
                {
                    "Variables": [f"{k[0]} vs {k[1]}" for k in partial_corr_results],
                    "Corrélation_Partielle": list(partial_corr_results.values()),
                }
            )
            partial_corr_df.to_excel(writer, sheet_name="Corrélations_Partielles", index=False)

        if include_stat_tests and klein_test_results:
            klein_df = pd.DataFrame(
                {
                    "Variables": [f"{k[0]} vs {k[1]}" for k in klein_test_results],
                    "R²_modèle": [v[0] for v in klein_test_results.values()],
                    "r²_variables": [v[1] for v in klein_test_results.values()],
                    "Seuil alpha": [alpha] * len(klein_test_results),
                    "Problème": [
                        "Oui" if v[0] < v[1] else "Non" for v in klein_test_results.values()
                    ],
                }
            )
            klein_df.to_excel(writer, sheet_name="Test_Klein", index=False)

        if include_stat_tests and fg_test_results:
            fg_global_df = pd.DataFrame(
                {
                    "Test": ["Chi² global"],
                    "Statistique": [fg_test_results["global"][0]],
                    "Degrés liberté": [fg_test_results["global"][1]],
                    "p-value": [fg_test_results["global"][2]],
                    "Seuil alpha": [alpha],
                    "Conclusion": [
                        "Multicolinéarité"
                        if fg_test_results["global"][2] < alpha
                        else "Pas de multicolinéarité"
                    ],
                }
            )
            fg_global_df.to_excel(writer, sheet_name="Test_FG_Global", index=False)

            fg_f_df = pd.DataFrame(
                {
                    "Variable": list(fg_test_results["f_tests"].keys()),
                    "F-statistique": [v[0] for v in fg_test_results["f_tests"].values()],
                    "p-value": [v[1] for v in fg_test_results["f_tests"].values()],
                    "Seuil alpha": [alpha] * len(fg_test_results["f_tests"]),
                    "Conclusion": [
                        "Liée aux autres" if v[1] < alpha else "Non liée"
                        for v in fg_test_results["f_tests"].values()
                    ],
                }
            )
            fg_f_df.to_excel(writer, sheet_name="Test_FG_F", index=False)

            if "t_tests" in fg_test_results:
                fg_t_df = pd.DataFrame(
                    {
                        "Variables": [f"{k[0]} vs {k[1]}" for k in fg_test_results["t_tests"]],
                        "t-statistique": [v[0] for v in fg_test_results["t_tests"].values()],
                        "p-value": [v[1] for v in fg_test_results["t_tests"].values()],
                        "Seuil alpha": [alpha] * len(fg_test_results["t_tests"]),
                        "Conclusion": [
                            "Corrélation" if v[1] < alpha else "Pas de corrélation"
                            for v in fg_test_results["t_tests"].values()
                        ],
                    }
                )
                fg_t_df.to_excel(writer, sheet_name="Test_FG_t", index=False)

        if include_stat_tests and x_var_names is not None:
            try:
                y_mean = np.mean(y)
                x_means = [np.mean(data[var]) for var in x_var_names]
                coefficients = [model.params[var] for var in x_var_names]

                elasticities = [
                    coef * (x_mean / y_mean) for coef, x_mean in zip(coefficients, x_means)
                ]
                sum_elasticities = sum(abs(e) for e in elasticities)
                optimal_shares = [abs(e) / sum_elasticities for e in elasticities]

                mix_df = pd.DataFrame(
                    {
                        "Variable": x_var_names,
                        "Elasticité": elasticities,
                        "Part_optimale": optimal_shares,
                    }
                )
                mix_df.to_excel(writer, sheet_name="Mix_optimal", index=False)
            except (KeyError, ZeroDivisionError):
                pass


def export_results_to_word(
    file_path: str,
    model: RegressionResultsWrapper,
    data: pd.DataFrame,
    X: pd.DataFrame,
    y: pd.Series,
    alpha: float = 0.05,
    decimal_places: int = 4,
    include_data: bool = True,
    include_summary: bool = True,
    include_diagnostics: bool = False,
    include_partial_corr: bool = False,
    include_stat_tests: bool = False,
    include_interpretations: bool = False,
    include_formulas: bool = False,
    include_calculation_steps: bool = False,
    include_hypothesis: bool = False,
    partial_corr_results: dict[tuple[str, str], float] | None = None,
    klein_test_results: dict[tuple[str, str], tuple[float, float]] | None = None,
    fg_test_results: dict[str, Any] | None = None,
    x_var_names: list[str] | None = None,
    interpretations: str | None = None,
    formulas: str | None = None,
    calculation_steps: str | None = None,
) -> None:
    """Exporte les resultats de l'analyse econometrique vers Word."""
    import docx
    from statsmodels.stats.diagnostic import het_breuschpagan, normal_ad
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    from statsmodels.stats.stattools import durbin_watson, jarque_bera

    doc = docx.Document()

    doc.add_heading("Résultats Économétriques", 0)
    doc.add_paragraph("\n")

    if include_data:
        doc.add_heading("Données", level=1)
        doc.add_paragraph(f"Nombre d'observations: {len(data)}")
        doc.add_paragraph(f"Variables: {', '.join(data.columns)}")
        doc.add_paragraph("\n")

    if include_summary:
        doc.add_heading("Résumé du Modèle", level=1)

        stats = [
            ("Variable dépendante", model.model.endog_names),
            ("R²", f"{model.rsquared:.{decimal_places}f}"),
            ("R² ajusté", f"{model.rsquared_adj:.{decimal_places}f}"),
            ("F-statistique", f"{model.fvalue:.{decimal_places}f}"),
            ("Prob(F-statistique)", f"{model.f_pvalue:.{decimal_places}f}"),
            ("Log-vraisemblance", f"{model.llf:.{decimal_places}f}"),
            ("AIC", f"{model.aic:.{decimal_places}f}"),
            ("BIC", f"{model.bic:.{decimal_places}f}"),
            ("Nb. observations", str(model.nobs)),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "LightShading"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Statistique"
        hdr_cells[1].text = "Valeur"

        for stat, val in stats:
            row_cells = table.add_row().cells
            row_cells[0].text = stat
            row_cells[1].text = val

        doc.add_paragraph("\n")

        doc.add_heading("Coefficients", level=2)
        coef_table = doc.add_table(rows=1, cols=5)
        coef_table.style = "LightShading"
        hdr_cells = coef_table.rows[0].cells
        for i, header in enumerate(["Variable", "Coefficient", "Ecart-type", "t", "P>|t|"]):
            hdr_cells[i].text = header

        for i, var in enumerate(model.params.index):
            row_cells = coef_table.add_row().cells
            values = [
                var,
                f"{model.params.iloc[i]:.{decimal_places}f}",
                f"{model.bse.iloc[i]:.{decimal_places}f}",
                f"{model.tvalues.iloc[i]:.{decimal_places}f}",
                f"{model.pvalues.iloc[i]:.{decimal_places}f}",
            ]
            for j, val in enumerate(values):
                row_cells[j].text = val

        doc.add_paragraph("\n")

    if include_diagnostics:
        doc.add_heading("Diagnostics", level=1)
        doc.add_heading("Tests de Diagnostic", level=2)
        resid = model.resid
        exog = model.model.exog

        tests = [
            ("Durbin-Watson (autocorrélation)", f"{durbin_watson(resid):.{decimal_places}f}"),
            (
                "Breusch-Pagan (hétéroscédasticité)",
                f"LM = {het_breuschpagan(resid, exog)[0]:.{decimal_places}f}, "
                f"p-value = {het_breuschpagan(resid, exog)[1]:.{decimal_places}f}",
            ),
            (
                "Anderson-Darling (normalité)",
                f"A² = {normal_ad(resid)[0]:.{decimal_places}f}, "
                f"p-value = {normal_ad(resid)[1]:.{decimal_places}f}",
            ),
            (
                "Jarque-Bera (normalité)",
                f"JB = {jarque_bera(resid)[0]:.{decimal_places}f}, "
                f"p-value = {jarque_bera(resid)[1]:.{decimal_places}f}",
            ),
        ]

        for test, result in tests:
            doc.add_paragraph(f"{test}: {result}")

        doc.add_paragraph("\n")

    doc.add_heading("Multicolinéarité", level=1)
    doc.add_heading("Facteurs d'Inflation de la Variance (VIF)", level=2)
    vif_table = doc.add_table(rows=1, cols=2)
    vif_table.style = "LightShading"
    vif_table.rows[0].cells[0].text = "Variable"
    vif_table.rows[0].cells[1].text = "VIF"

    vif_data: list[tuple[str, float]] = []
    for i, col in enumerate(X.columns):
        if col == "const":
            continue
        vif = variance_inflation_factor(X.values, i)
        vif_data.append((col, vif))

    vif_data.sort(key=lambda x: x[1], reverse=True)

    for var, vif in vif_data:
        row_cells = vif_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = f"{vif:.{decimal_places}f}"

    doc.add_paragraph("\n")

    if include_stat_tests and klein_test_results:
        doc.add_heading("Test de Klein", level=2)
        klein_table = doc.add_table(rows=1, cols=4)
        klein_table.style = "LightShading"
        hdr_cells = klein_table.rows[0].cells
        hdr_cells[0].text = "Variables"
        hdr_cells[1].text = "R² modèle"
        hdr_cells[2].text = "r² variables"
        hdr_cells[3].text = "Problème"

        for (var1, var2), (r2, r2_vars) in klein_test_results.items():
            row_cells = klein_table.add_row().cells
            row_cells[0].text = f"{var1} vs {var2}"
            row_cells[1].text = f"{r2:.4f}"
            row_cells[2].text = f"{r2_vars:.4f}"
            row_cells[3].text = "Oui" if r2 < r2_vars else "Non"

        doc.add_paragraph("\n")

    if include_stat_tests and fg_test_results:
        doc.add_heading("Test de Farrar-Glauber", level=2)

        doc.add_heading("Test global du Chi²", level=3)
        fg_global_table = doc.add_table(rows=1, cols=4)
        fg_global_table.style = "LightShading"
        hdr_cells = fg_global_table.rows[0].cells
        hdr_cells[0].text = "Test"
        hdr_cells[1].text = "Statistique"
        hdr_cells[2].text = "p-value"
        hdr_cells[3].text = "Conclusion"

        row_cells = fg_global_table.add_row().cells
        row_cells[0].text = "Chi² global"
        row_cells[1].text = f"{fg_test_results['global'][0]:.4f}"
        row_cells[2].text = f"{fg_test_results['global'][2]:.4f}"
        row_cells[3].text = (
            "Multicolinéarité" if fg_test_results["global"][2] < 0.05 else "Pas de multicolinéarité"
        )
        doc.add_paragraph("\n")

        doc.add_heading("Tests F sur régressions auxiliaires", level=3)
        fg_f_table = doc.add_table(rows=1, cols=4)
        fg_f_table.style = "LightShading"
        hdr_cells = fg_f_table.rows[0].cells
        hdr_cells[0].text = "Variable"
        hdr_cells[1].text = "F-statistique"
        hdr_cells[2].text = "p-value"
        hdr_cells[3].text = "Conclusion"

        for var, (f_val, f_pval) in fg_test_results["f_tests"].items():
            row_cells = fg_f_table.add_row().cells
            row_cells[0].text = var
            row_cells[1].text = f"{f_val:.4f}"
            row_cells[2].text = f"{f_pval:.4f}"
            row_cells[3].text = "Liée aux autres" if f_pval < 0.05 else "Non liée"

        doc.add_paragraph("\n")

        if "t_tests" in fg_test_results:
            doc.add_heading("Tests t sur corrélations partielles", level=3)
            fg_t_table = doc.add_table(rows=1, cols=4)
            fg_t_table.style = "LightShading"
            hdr_cells = fg_t_table.rows[0].cells
            hdr_cells[0].text = "Variables"
            hdr_cells[1].text = "t-statistique"
            hdr_cells[2].text = "p-value"
            hdr_cells[3].text = "Conclusion"

            for (var1, var2), (t_val, t_pval) in fg_test_results["t_tests"].items():
                row_cells = fg_t_table.add_row().cells
                row_cells[0].text = f"{var1} vs {var2}"
                row_cells[1].text = f"{t_val:.4f}"
                row_cells[2].text = f"{t_pval:.4f}"
                row_cells[3].text = "Corrélation" if t_pval < 0.05 else "Pas de corrélation"

            doc.add_paragraph("\n")

    if include_partial_corr and partial_corr_results:
        doc.add_heading("Corrélations Partielles", level=1)
        partial_table = doc.add_table(rows=1, cols=3)
        partial_table.style = "LightShading"
        hdr_cells = partial_table.rows[0].cells
        hdr_cells[0].text = "Variable 1"
        hdr_cells[1].text = "Variable 2"
        hdr_cells[2].text = "Corrélation Partielle"

        for (var1, var2), corr in partial_corr_results.items():
            row_cells = partial_table.add_row().cells
            row_cells[0].text = var1
            row_cells[1].text = var2
            row_cells[2].text = f"{corr:.{decimal_places}f}"

        doc.add_paragraph("\n")

    if include_stat_tests and x_var_names is not None:
        try:
            y_mean = np.mean(y)
            x_means = [np.mean(data[var]) for var in x_var_names]
            coefficients = [model.params[var] for var in x_var_names]

            elasticities = [coef * (x_mean / y_mean) for coef, x_mean in zip(coefficients, x_means)]
            sum_elasticities = sum(abs(e) for e in elasticities)
            optimal_shares = [abs(e) / sum_elasticities for e in elasticities]

            doc.add_heading("Mix Optimal", level=1)
            mix_table = doc.add_table(rows=1, cols=3)
            mix_table.style = "LightShading"
            hdr_cells = mix_table.rows[0].cells
            hdr_cells[0].text = "Variable"
            hdr_cells[1].text = "Elasticité"
            hdr_cells[2].text = "Part optimale"

            for var, elasticity, share in zip(x_var_names, elasticities, optimal_shares):
                row_cells = mix_table.add_row().cells
                row_cells[0].text = var
                row_cells[1].text = f"{elasticity:.4f}"
                row_cells[2].text = f"{share * 100:.2f}%"

            doc.add_paragraph("\n")
        except (KeyError, ZeroDivisionError):
            pass

    if include_interpretations and interpretations:
        doc.add_heading("Interprétations", level=1)
        for interp in interpretations.split("\n"):
            if interp.strip():
                doc.add_paragraph(interp)
        doc.add_paragraph("\n")

    if include_formulas and formulas:
        doc.add_heading("Formules Mathématiques", level=1)
        for formula in formulas.split("\n"):
            if formula.strip():
                p = doc.add_paragraph()
                p.add_run(formula).font.name = "Courier New"
        doc.add_paragraph("\n")

    if include_calculation_steps and calculation_steps:
        doc.add_heading("Étapes de Calcul", level=1)
        for step in calculation_steps.split("\n"):
            if step.strip():
                p = doc.add_paragraph()
                p.add_run(step).font.name = "Courier New"
        doc.add_paragraph("\n")

    if include_hypothesis:
        doc.add_heading("Hypothèses des Tests", level=1)

        doc.add_heading("Test t de significativité", level=2)
        doc.add_paragraph(
            "H0: Le coefficient est égal à 0\n"
            "H1: Le coefficient est différent de 0\n"
            "Règle de décision: Rejet de H0 si p-value < 0.05"
        )
        doc.add_paragraph("\n")

        doc.add_heading("Test F global", level=2)
        doc.add_paragraph(
            "H0: Tous les coefficients (sauf constante) = 0\n"
            "H1: Au moins un coefficient ≠ 0\n"
            "Règle de décision: Rejet de H0 si p-value < 0.05"
        )
        doc.add_paragraph("\n")

        doc.add_heading("Test de Klein", level=2)
        doc.add_paragraph(
            "H0: Pas de multicolinéarité problématique\n"
            "H1: Présence de multicolinéarité problématique\n"
            "Règle: Si R² modèle < r² entre deux variables ⇒ multicolinéarité"
        )
        doc.add_paragraph("\n")

        doc.add_heading("Test de Farrar-Glauber", level=2)
        doc.add_paragraph(
            "1. Test global du Chi²:\n"
            "   H0: Matrice de corrélation = Matrice identité\n"
            "   H1: Matrice de corrélation ≠ Matrice identité\n"
            "   Règle: Rejet de H0 si p-value < 0.05\n\n"
            "2. Tests F sur régressions auxiliaires:\n"
            "   H0: Tous les coefficients = 0\n"
            "   H1: Au moins un coefficient ≠ 0\n"
            "   Règle: Rejet de H0 si p-value < 0.05\n\n"
            "3. Tests t sur corrélations partielles:\n"
            "   H0: ρ(Xi,Xj|autres) = 0\n"
            "   H1: ρ(Xi,Xj|autres) ≠ 0\n"
            "   Règle: Rejet de H0 si p-value < 0.05"
        )
        doc.add_paragraph("\n")

    doc.save(file_path)


def export_results_to_pdf(
    file_path: str,
    model: RegressionResultsWrapper,
    data: pd.DataFrame,
    X: pd.DataFrame,
    y: pd.Series,
    alpha: float = 0.05,
    decimal_places: int = 4,
    include_data: bool = True,
    include_summary: bool = True,
    include_diagnostics: bool = False,
    include_partial_corr: bool = False,
    include_stat_tests: bool = False,
    partial_corr_results: dict[tuple[str, str], float] | None = None,
    klein_test_results: dict[tuple[str, str], tuple[float, float]] | None = None,
    fg_test_results: dict[str, Any] | None = None,
    x_var_names: list[str] | None = None,
) -> None:
    """Exporte les resultats de l'analyse econometrique vers PDF (format paysage)."""
    from fpdf import FPDF
    from statsmodels.stats.diagnostic import het_breuschpagan, normal_ad
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    from statsmodels.stats.stattools import durbin_watson, jarque_bera

    pdf = FPDF(orientation="L")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=10)

    usable_width = 280

    pdf.set_font("Arial", "B", 14)
    pdf.cell(usable_width, 10, txt="Résultats Économétriques", ln=1, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(
        usable_width,
        8,
        txt=f"Généré le {__import__('datetime').datetime.now().strftime('%d/%m/%Y à %H:%M')}",
        ln=1,
        align="C",
    )
    pdf.ln(5)

    line_height = 8

    if include_data and hasattr(data, "columns"):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(usable_width, 8, txt="Données", ln=1)
        pdf.set_font("Arial", "", 10)

        col_widths = [60, 220]
        pdf.cell(col_widths[0], 8, txt="Nombre d'observations:", border=1)
        pdf.cell(col_widths[1], 8, txt=str(len(data)), border=1, ln=1)

        if hasattr(data, "columns"):
            pdf.cell(col_widths[0], 8, txt="Variables:", border=1)
            vars_text = ", ".join(data.columns)
            if len(vars_text) > 100:
                vars_text = vars_text[:100] + "..."
            pdf.cell(col_widths[1], 8, txt=vars_text, border=1, ln=1)

        pdf.ln(5)

    if include_summary and hasattr(model, "model"):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(usable_width, 8, txt="Résumé du Modèle", ln=1)
        pdf.set_font("Arial", "", 10)

        stats = [
            ("Variable dépendante", getattr(model.model, "endog_names", "N/A")),
            ("R²", f"{getattr(model, 'rsquared', 0):.{decimal_places}f}"),
            ("R² ajusté", f"{getattr(model, 'rsquared_adj', 0):.{decimal_places}f}"),
            ("F-statistique", f"{getattr(model, 'fvalue', 0):.{decimal_places}f}"),
            ("Prob(F-statistique)", f"{getattr(model, 'f_pvalue', 0):.{decimal_places}f}"),
            ("Log-vraisemblance", f"{getattr(model, 'llf', 0):.{decimal_places}f}"),
            ("AIC", f"{getattr(model, 'aic', 0):.{decimal_places}f}"),
            ("BIC", f"{getattr(model, 'bic', 0):.{decimal_places}f}"),
            ("Nb. observations", str(getattr(model, "nobs", 0))),
        ]

        col_width = usable_width / 2
        for stat, val in stats:
            pdf.cell(col_width, line_height, txt=stat, border=1)
            pdf.cell(col_width, line_height, txt=str(val), border=1, ln=1)

        pdf.ln(5)

        if hasattr(model, "params"):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(usable_width, 8, txt="Coefficients", ln=1)
            pdf.set_font("Arial", "", 10)

            headers = ["Variable", "Coefficient", "Ecart-type", "t", "P>|t|"]
            c_widths = [
                usable_width * 0.25,
                usable_width * 0.20,
                usable_width * 0.20,
                usable_width * 0.15,
                usable_width * 0.20,
            ]

            for i, header in enumerate(headers):
                pdf.cell(c_widths[i], line_height, txt=header, border=1, align="C")
            pdf.ln()

            for i, var in enumerate(model.params.index):
                values = [
                    str(var),
                    f"{model.params.iloc[i]:.{decimal_places}f}",
                    f"{getattr(model, 'bse', [0] * len(model.params)).iloc[i]:.{decimal_places}f}",
                    f"{getattr(model, 'tvalues', [0] * len(model.params)).iloc[i]:.{decimal_places}f}",
                    f"{getattr(model, 'pvalues', [0] * len(model.params)).iloc[i]:.{decimal_places}f}",
                ]
                for j, val in enumerate(values):
                    align = "R" if j > 0 else "L"
                    pdf.cell(c_widths[j], line_height, txt=val, border=1, align=align)
                pdf.ln()

            pdf.ln(5)

    if include_diagnostics and hasattr(model, "resid"):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(usable_width, 8, txt="Diagnostics", ln=1)
        pdf.set_font("Arial", "", 10)

        resid = model.resid
        exog = model.model.exog

        try:
            dw = durbin_watson(resid)
        except Exception:
            dw = "N/A"

        try:
            bp_test = het_breuschpagan(resid, exog)
        except Exception:
            bp_test = ("N/A", "N/A")

        try:
            ad_test = normal_ad(resid)
        except Exception:
            ad_test = ("N/A", "N/A")

        try:
            jb_test = jarque_bera(resid)
        except Exception:
            jb_test = ("N/A", "N/A")

        tests = [
            (
                "Durbin-Watson (autocorrélation)",
                f"{dw:.{decimal_places}f}" if isinstance(dw, (int, float)) else str(dw),
            ),
            (
                "Breusch-Pagan (hétéroscédasticité)",
                f"LM={bp_test[0]:.{decimal_places}f}, p={bp_test[1]:.{decimal_places}f}"
                if all(isinstance(x, (int, float)) for x in bp_test)
                else "N/A",
            ),
            (
                "Anderson-Darling (normalité)",
                f"A²={ad_test[0]:.{decimal_places}f}, p={ad_test[1]:.{decimal_places}f}"
                if all(isinstance(x, (int, float)) for x in ad_test)
                else "N/A",
            ),
            (
                "Jarque-Bera (normalité)",
                f"JB={jb_test[0]:.{decimal_places}f}, p={jb_test[1]:.{decimal_places}f}"
                if all(isinstance(x, (int, float)) for x in jb_test)
                else "N/A",
            ),
        ]

        col_width = usable_width / 2
        for i in range(0, len(tests), 2):
            test1, result1 = tests[i]
            pdf.cell(col_width, 8, txt=test1, border=1)
            pdf.cell(col_width, 8, txt=result1, border=1, ln=1)

            if i + 1 < len(tests):
                test2, result2 = tests[i + 1]
                pdf.cell(col_width, 8, txt=test2, border=1)
                pdf.cell(col_width, 8, txt=result2, border=1, ln=1)

        pdf.ln(5)

    if hasattr(X, "columns"):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(usable_width, 8, txt="Multicolinéarité", ln=1)
        pdf.set_font("Arial", "", 10)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(usable_width, 8, txt="Facteurs d'Inflation de la Variance (VIF)", ln=1)
        pdf.set_font("Arial", "", 10)

        vif_data: list[tuple[str, float]] = []
        for i, col in enumerate(X.columns):
            if col == "const":
                continue
            try:
                vif = variance_inflation_factor(X.values, i)
                vif_data.append((col, vif))
            except Exception:
                vif_data.append((col, "N/A"))

        if vif_data:
            vif_data.sort(
                key=lambda x: x[1] if isinstance(x[1], (int, float)) else 0,
                reverse=True,
            )

            col_width = usable_width / 2
            for i in range(0, len(vif_data), 2):
                var1, vif1 = vif_data[i]
                pdf.cell(col_width * 0.6, line_height, txt=var1, border=1)
                pdf.cell(
                    col_width * 0.4,
                    line_height,
                    txt=f"{vif1:.{decimal_places}f}"
                    if isinstance(vif1, (int, float))
                    else str(vif1),
                    border=1,
                )

                if i + 1 < len(vif_data):
                    var2, vif2 = vif_data[i + 1]
                    pdf.cell(col_width * 0.6, line_height, txt=var2, border=1)
                    pdf.cell(
                        col_width * 0.4,
                        line_height,
                        txt=f"{vif2:.{decimal_places}f}"
                        if isinstance(vif2, (int, float))
                        else str(vif2),
                        border=1,
                    )

                pdf.ln()

            pdf.ln(5)

    pdf.output(file_path)
